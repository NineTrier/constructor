import json
import base64
import re

from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic import CreateView
from .models import DocumentsPattern, DocumentPattern_Objects,Fonts, SavedElements, VariableBlock, DocType, Document_ParentDocument, Document_VariableBlock
from database_manager.models import Object, Parameter
from .forms import DocumentForm
from django.core.files.storage import FileSystemStorage
import os
from django.conf import settings
from django.http import HttpRequest, HttpResponse, Http404, HttpResponseNotModified, HttpResponseForbidden
from docx import Document
from .Document import Document

from django.views.decorators.csrf import csrf_exempt
from transliterate import translit
from transliterate.decorators import transliterate_function
from user_manager.models import Profile, user_directory_path

import pymorphy3
from pymorphy3.shapes import restore_capitalization

raskrit = {
    'ООО': 'Общество с ограниченной ответственностью',
    'ИП':'Индивидуальный предприниматель',
    'ПАО': 'Публичное акционерное общество',
    'АО': 'Акционерное общество'
}

matchers ={ 
    r'\s?Арбитражн\w*\sсуд\w?': r'\s?Арбитражн\w*\sсуд\w?\s',
    r'\s?Обществ\w*\sс\sограниченной\sответственностью\w*': r'\s?Обществ\w+?\s',
    r'\s?Индивидуальн\w*\sпредпринимател\w?': r'\s?Индивидуальн\w*\sпредпринимател\w?\s',
    r'\s?Акционерн\w*\sобществ\w?': r'\s?Акционерн\w*\sобществ\w?\s',
}

morph = pymorphy3.MorphAnalyzer()

# Класс, который помогает создавать новый записи в базу данных Documents
# и открывает страницу с добавлением новых документов
class DocumentCreate(CreateView):
    model = DocumentsPattern
    form_class = DocumentForm

    template_name = 'document/document_create.html'
    
    def get_context_data(self, **kwargs):
        context = super(DocumentCreate, self).get_context_data(**kwargs)
        if self.request.user.is_authenticated:
            context['profile'] = Profile.objects.filter(user=self.request.user)[0]
        context['documents'] = DocumentsPattern.objects.all()
        context['title'] = "Загрузка документа"
        return context
    
    def post(self, request, *args, **kwargs):
        document = DocumentsPattern()
        document.name = request.POST['name']
        document.description = request.POST['description']
        document.type = DocType.objects.get(id=request.POST['type'])
        document.owner = Profile.objects.filter(user=request.user)[0]
        file = request.FILES.get('file')
        fileDirect = int(DocumentsPattern.objects.latest('id').id)+1
        if not os.path.isdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}"):
            os.mkdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}")
        if not os.path.isdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}"):
            os.mkdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}")
        fs = FileSystemStorage(location=f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}")
        fs.save(translit_russian(file.name), file)
        file_url = f"documents\\user_{request.user.id}\\{fileDirect}\\{translit_russian(file.name)}"
        document.file = file_url
        document.save()
        response = redirect(f"/document/view?id={document.id}&type=1")
        return response

def translit_russian(text):
    """Выполняет транслитерацию русского текста в латиницу для названий файлов"""
    try:
        translited = translit(text, reversed=True)
        translited = str(translited).replace(' ', '_')
        return translited
    except Exception as exc:
        print(exc)
        return text
    

def create_New_Document(request):
    """Обработчик запроса для создания документа"""
    if request.method == 'POST':
        document = DocumentsPattern()
        document.name = request.POST['name']
        document.description = request.POST['description']
        document.type = DocType.objects.get(id=request.POST['type'])
        document.owner = Profile.objects.filter(user=request.user)[0]
        document.picture = f"noimage.jpeg"
        print(document)
        try:
            fileDirect = int(DocumentsPattern.objects.latest('id').id)+1
        except Exception as exc:
            print(exc)
            fileDirect = 0
        print(fileDirect)
        os.makedirs(os.path.dirname(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}"), exist_ok=True)
        os.makedirs(os.path.dirname(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}"), exist_ok=True)
        file_name = f"{translit_russian(document.name)}{translit_russian(document.owner.__str__())}"[:50]
        file_url = f"documents\\user_{request.user.id}\\{fileDirect}\\{file_name}.docx"
        document.file = file_url
        doc_file = Document(f"{settings.MEDIA_ROOT}\\blank.docx")
        os.makedirs(os.path.dirname(f"{settings.MEDIA_ROOT}\\{file_url}"), exist_ok=True)
        doc_file.save(f"{settings.MEDIA_ROOT}\\{file_url}")
        document.save()
        return redirect(f"/document/view?id={document.id}&type=1")
    form = DocumentForm()
    context = {
        'form': form,
        'title': 'Создание документа'
    }
    if request.user.is_authenticated:
        context['profile'] = Profile.objects.filter(user=request.user)[0]
    return render(request, 'document/new_document.html', context)

@csrf_exempt
def SaveCover(request):
    """Обработчик запроса для загрузки обложки"""
    try:
        id = request.POST['id']
        document = DocumentsPattern.objects.get(id=id)
        img = request.POST['img']
        img = str(img).replace('data:image/png;base64,', '')
        img = str(img).replace(' ', '+')
        dat = base64.decodebytes(img.encode('utf-8'))
        fileDirect = document.id
        if not os.path.isdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}"):
            os.mkdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}")
        with open(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}\\{translit_russian(document.name)}Cover.png", 'wb') as file:
            file.write(dat)
        file_url = f"documents\\user_{request.user.id}\\{fileDirect}\\{translit_russian(document.name)}Cover.png"
        document.picture = file_url
        document.save() 
        return HttpResponse()
    except Exception as exc:
        print(exc)
        return HttpResponseNotModified()

@csrf_exempt
def SaveImage(request):
    """Обработчик запроса для загрузки картинок в документ"""
    try:
        id = request.POST['id']
        document = DocumentsPattern.objects.get(id=id)
        img = request.FILES['img']
        fileDirect = document.id
        if not os.path.isdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}"):
            os.mkdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}")
        file_url = f"documents\\user_{request.user.id}\\{fileDirect}\\{request.POST['idImage']}.png"
        if os.path.isfile(f"{settings.MEDIA_ROOT}\\{file_url}"):
            os.remove(f"{settings.MEDIA_ROOT}\\{file_url}")
        fs = FileSystemStorage()
        filename = fs.save(file_url, img)
        response = HttpResponse()
        response['img'] = f"{settings.MEDIA_URL}{file_url}"
        return response
    except Exception as exc:
        print(exc)
        return HttpResponseNotModified()

# Функция, которая позволяет скачать файл, загруженный на сервер
def download(request):
    """Обработчик запроса для выгрузки документа"""
    document = DocumentsPattern.objects.filter(id=request.GET.get('id'))[0]
    document.downloadsTimes = int(document.downloadsTimes) + 1
    document.save()
    filepath = document.file
    file_path = os.path.join(settings.MEDIA_ROOT, str(filepath))
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
            response['Content-Disposition'] = 'inline; filename=' + f'{translit_russian(document.name)}.docx'
            return response
    raise Http404

@csrf_exempt
def DeleteSavedElement(request):
    """Обработчик запроса для удаления сохраненного элемента"""
    response = redirect('/')
    request_data = request.body
    stroke = json.loads(request_data)
    if stroke['id'] == '-1':
        return response
    SavedElements.objects.filter(id=stroke['id']).delete()
    return response

@csrf_exempt
def DeleteVariable(request):
    """Обработчик запроса для удаления переменной"""
    response = HttpResponse()
    request_data = request.body
    stroke = json.loads(request_data)
    if stroke['id'] == '-1':
        return response
    VariableBlock.objects.filter(id=stroke['id']).delete()
    return response

@csrf_exempt
def CreateDocType(request):
    """Обработчик запроса для создания типа документа"""
    try:
        response = HttpResponse()
        request_data = request.body
        stroke = json.loads(request_data)
        type = DocType()
        type.name = stroke['name']
        type.save()
        response['id'] = type.id
        return response
    except Exception as exc:
        print(exc)
        return HttpResponseNotModified()

@csrf_exempt
def DeleteDocType(request):
    """Обработчик запроса для удаления типа документа"""
    try:
        response = HttpResponse()
        request_data = request.body
        stroke = json.loads(request_data)
        type = DocType.objects.get(id=stroke['id'])
        type.delete()
        return response
    except Exception as exc:
        print(exc)
        return HttpResponseNotModified()

@csrf_exempt
def SaveVariable(request):
    """Обработчик запроса для сохранения переменной в базу данных"""
    request_data = request.body
    stroke = json.loads(request_data)
    doc_variables = Document_VariableBlock.objects.filter(document=stroke['id_doc'])
    if doc_variables == None:
        doc_variables = []
    for doc_variable in doc_variables:
        if doc_variable.variable.name == stroke['name']:
            response = HttpResponseNotModified()
            response['MessageOfError'] = 'Переменная с таким именем существует.'.encode('utf-8')
            response['TypeOfError'] = 'Ошибка переменных'.encode('utf-8')
            return response
    response = HttpResponse()
    if stroke['id'] == '-1':
        variable = VariableBlock()
    else:
        variable = VariableBlock.objects.filter(id=stroke['id'])[0]
    variable.name = stroke['name']
    variable.meaning = stroke['value']
    variable.save()
    new_doc_variable = Document_VariableBlock()
    new_doc_variable.document = DocumentsPattern.objects.filter(id=stroke['id_doc'])[0]
    new_doc_variable.variable = variable
    new_doc_variable.save()
    response['id'] = variable.id
    return response

@csrf_exempt
def CopyDocument(request, id=None):
    """Обработчик запроса для копирования документа пользователю"""
    try:
        response = HttpResponse()
        request_data = request.body
        if id == None:
            stroke = json.loads(request_data)
        else:
            stroke = {'id': id}
        print(stroke)
        profile = Profile.objects.filter(user=request.user)[0]
        documentToCopy = DocumentsPattern.objects.get(id=stroke['id'])
        objects = DocumentPattern_Objects.objects.filter(document=documentToCopy)
        variables = Document_VariableBlock.objects.filter(document=documentToCopy.id)
        document = DocumentsPattern()
        document.name = documentToCopy.name + '-копия'
        document.owner = profile
        document.type = documentToCopy.type
        document.description = documentToCopy.description
        document.picture = f"noimage.jpeg"
        document.json = documentToCopy.json
        fileDirect = int(DocumentsPattern.objects.latest('id').id)+1
        if not os.path.isdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}"):
            os.mkdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}")
        if not os.path.isdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}"):
            os.mkdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}")
        file_name = f"{translit_russian(document.name)}{translit_russian(document.owner.__str__())}.docx"
        file_url = f"documents\\user_{request.user.id}\\{fileDirect}\\{file_name}"
        doc_file = Document(documentToCopy.file)
        doc_file.save(f"{settings.MEDIA_ROOT}/{file_url}")
        document.file = file_url
        document.documentOfOrganisation = False
        if document.save():
            print('Документ добавлен')
            response['id'] = document.id
            for document_variable in variables:
                print(document_variable)
                new_document_variable = Document_VariableBlock.objects.filter(document=document, variable=document_variable.variable).first()
                if not new_document_variable:
                    new_document_variable = Document_VariableBlock()
                new_document_variable = Document_VariableBlock()
                new_document_variable.variable = document_variable.variable
                new_document_variable.document = document
                new_document_variable.save()
            for document_object in objects:
                new_document_object = DocumentPattern_Objects.objects.filter(object=document_object.object, document=document).first()
                if not new_document_object:
                    new_document_object = DocumentPattern_Objects()
                new_document_object.document = document
                new_document_object.object = document_object.object
                new_document_object.save()
            return response
        else:
            print('документ не сохранен')
            return HttpResponseNotModified()
    except Exception as exc:
        print(f"Не получилось {exc}")
        return HttpResponseNotModified()

@csrf_exempt
def AddDocumentToUser(request, id=None):
    """Обработчик запроса для копирования документа пользователю"""
    try:
        response = HttpResponse()
        request_data = request.body
        if id == None:
            stroke = json.loads(request_data)
        else:
            stroke = {'id': id}
        profile = Profile.objects.filter(user=request.user)[0]
        documentToCopy = DocumentsPattern.objects.get(id=stroke['id'])
        document_parentDocument = Document_ParentDocument.objects.filter(parent=documentToCopy).filter(userRequested=profile)
        objects = DocumentPattern_Objects.objects.filter(document=documentToCopy)
        variables = Document_VariableBlock.objects.filter(document=documentToCopy.id)
        if document_parentDocument:
            document = DocumentsPattern.objects.get(id=document_parentDocument[0].document.id)
            document_parentDocument = document_parentDocument[0]
        else:
            document = DocumentsPattern()
            document_parentDocument = Document_ParentDocument()
            document_parentDocument.document = document
            document_parentDocument.parent = documentToCopy
            document_parentDocument.userRequested = profile
            document_parentDocument.parentDocumentChanged = False
        document.name = documentToCopy.name[:40]
        document.owner = profile
        document.type = documentToCopy.type
        document.description = documentToCopy.description
        document.picture = f"noimage.jpeg"
        document.json = documentToCopy.json
        fileDirect = int(DocumentsPattern.objects.latest('id').id)+1
        if not os.path.isdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}"):
            os.mkdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}")
        if not os.path.isdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}"):
            os.mkdir(f"{settings.MEDIA_ROOT}\\documents\\user_{request.user.id}\\{fileDirect}")
        file_name = f"{translit_russian(document.name)}{translit_russian(document.owner.__str__())}.docx"
        file_url = f"documents\\user_{request.user.id}\\{fileDirect}\\{file_name}"
        doc_file = Document(documentToCopy.file)
        doc_file.save(f"{settings.MEDIA_ROOT}/{file_url}")
        document.file = file_url
        document.documentOfOrganisation = False
        if document.save():
            print('Документ добавлен')
            response['id'] = document.id
            document_parentDocument.save()
            for document_variable in variables:
                print(document_variable)
                new_document_variable = Document_VariableBlock.objects.filter(document=document, variable=document_variable.variable).first()
                if not new_document_variable:
                    new_document_variable = Document_VariableBlock()
                new_document_variable = Document_VariableBlock()
                new_document_variable.variable = document_variable.variable
                new_document_variable.document = document
                new_document_variable.save()
            for document_object in objects:
                new_document_object = DocumentPattern_Objects.objects.filter(object=document_object.object, document=document).first()
                if not new_document_object:
                    new_document_object = DocumentPattern_Objects()
                new_document_object.document = document
                new_document_object.object = document_object.object
                new_document_object.save()
            return response
        else:
            print('документ не сохранен')
            return HttpResponseNotModified()
    except Exception as exc:
        print(f"Не получилось {exc}")
        return HttpResponseNotModified()
    


@csrf_exempt
def SaveSavedElement(request):
    """Обработчик запроса для сохранения сохраненного элемента"""
    response = HttpResponse()
    request_data = request.body
    stroke = json.loads(request_data)
    if stroke['id'] == '-1':
        saved_element = SavedElements()
    else:
        saved_element = SavedElements.objects.filter(id=stroke['id'])[0]
    saved_element.name = stroke['name']
    saved_element.json = stroke['json']
    saved_element.save()
    response['id'] = saved_element.id
    return response

@csrf_exempt
def DeleteDocument(request):
    """Обработчик запроса для удаления документа из базы данных"""
    response = redirect('/')
    request_data = request.body
    stroke = json.loads(request_data) 
    print(stroke)
    if stroke['id'] == '-1':
        return response
    document = DocumentsPattern.objects.get(id=stroke['id'])
    print(document)
    childs = Document_ParentDocument.objects.filter(parent=document) | Document_ParentDocument.objects.filter(document=document)
    print(childs)
    for child in childs:
        child.delete()
    try:
        os.remove(f"{settings.MEDIA_ROOT}/{document.file}")
        os.remove(f"{settings.MEDIA_ROOT}/{document.picture}")
        document.delete()
        return response
    except Exception as exc:
        print(document)
        try:
            doc_deleted = document.delete()
        except Exception as exc:
            print(exc)
        print("#######", doc_deleted)
        print(exc)
        return response
    
def replace_last(source_string, replace_what, replace_with):
    head, _sep, tail = source_string.rpartition(replace_what)
    return head + replace_with + tail

def ViewDocument(request):
    """Обработчик запроса для просмотра документа"""
    fileid = request.GET.get('id')
    Doc = get_object_or_404(DocumentsPattern, pk=fileid)
    file_path = Doc.file
    file_path = os.path.join(settings.MEDIA_ROOT, str(file_path))
    if request.user != Doc.owner.user:
        res = AddDocumentToUser(request, fileid)
        return redirect(f'/document/view?id={res["id"]}')
    parent_document = Document_ParentDocument.objects.filter(document=Doc.id)
    objects = [{'object': obj.object, 'params':[parameter for parameter in sorted(Parameter.objects.filter(object=obj.object), key=lambda x: x.id)]} for obj in DocumentPattern_Objects.objects.filter(document=Doc.id)]
    print(objects)
    context = {
        'title': 'Просмотр документа',
        'Doc': Doc,
        'fonts': Fonts.objects.order_by('id'),
        'variable': json.dumps({var.variable.id: f"{var.variable.name}:{var.variable.meaning}" for var in Document_VariableBlock.objects.filter(document=Doc.id)}),
        'document_json': json.dumps(Doc.json),
        'objects': objects,
    }
    if request.user.is_authenticated:
        context['profile'] = Profile.objects.filter(user=request.user)[0]
    return render(request, 'document/document_view_v3.html', context)

def CreateDocumentMultiple(request):
    if request.method == 'GET':
        documents = request.POST.getlist('object_documents[]')
        idents = request.POST.getlist('object_idents[]')
        for doc_id in documents:
            return redirect(f'/document/view?id={res["id"]}')

def ViewDocumentAndCreateDocument(request):
    """Обработчик запроса для просмотра документа"""
    fileid = request.GET.get('id')
    
    Doc = get_object_or_404(DocumentsPattern, pk=fileid)
    file_path = Doc.file
    file_path = os.path.join(settings.MEDIA_ROOT, str(file_path))
    if request.user != Doc.owner.user:
        res = AddDocumentToUser(request, fileid)
        return redirect(f'/document/view?id={res["id"]}')
    parent_document = Document_ParentDocument.objects.filter(document=Doc.id)
    objects = [{'object': obj.object, 'params':[parameter for parameter in Parameter.objects.filter(object=obj.object)]} for obj in DocumentPattern_Objects.objects.filter(document=Doc.id)]
    print(objects)
    context = {
        'title': 'Просмотр документа',
        'Doc': Doc,
        'fonts': Fonts.objects.order_by('id'),
        'variable': json.dumps({var.variable.id: f"{var.variable.name}:{var.variable.meaning}" for var in Document_VariableBlock.objects.filter(document=Doc.id)}),
        'document_json': json.dumps(Doc.json),
        'objects': objects,
        'create_document': True,
    }
    if request.user.is_authenticated:
        context['profile'] = Profile.objects.filter(user=request.user)[0]
    return render(request, 'document/document_view_v3.html', context)

def SavedElementFromJSON(json_stroke):
    """Функция создаёт элементы класса SavedElements из JSON-строки"""
    for list_elem in json_stroke:
        name, id_elem, element = list_elem['name'], list_elem['id'], list_elem['json']
        if id_elem == '-1':
            saved = SavedElements()
            saved.name = name
            saved.json = json.loads(element)
            saved.save()
        else:
            saved = SavedElements.objects.get(id=id_elem)
            saved.name = name
            saved.json = json.loads(element)
            saved.save()

@csrf_exempt
def UpdateDocument(request):
    """Функция обрабатывает запрос и обновляет документ в базе данных"""
    fileid = request.GET.get('id')
    doc = DocumentsPattern.objects.get(id=fileid)
    childs = Document_ParentDocument.objects.filter(parent=doc)
    print(childs)
    for child in childs:
        print(child.document.id)
    file_path = os.path.join(settings.MEDIA_ROOT, str(doc.file))
    document = Document()
    request_data = request.body
    stroke = json.loads(request_data)
    doc.json = stroke
    doc.name = stroke['doc_name']
    if not doc.save():
        response = HttpResponseNotModified()
        response['MessageOfError'] = "Документ не сохранён. Попробуйте позже.".encode('utf-8')
        response['TypeOfError'] = "Документ не сохранён".encode('utf-8')
        return response
    document.from_json(stroke)
    document.save(file_path)
    response = HttpResponse()
    return response

@csrf_exempt
def AcceptFilters(request):
    request_data = request.body
    stroke = json.loads(request_data)
    print(stroke)
    result = {}
    for key, value in stroke.items():
        for filter in value['filters']:
            if filter == 'Raskrit':
                value['phrase'] = Raskritie(value)
            if filter == 'ChangeCattle':
                value['phrase'] = ChangeCattle(value)
            if filter == 'ChangeCase':
                value['phrase'] = UpperCase(value)
        result[key] = str(value['phrase']).replace('ё', 'е').replace('Ё', 'Е')
    response = HttpResponse()
    response.content = json.dumps(result)
    response.charset = 'utf-8'
    return response

def Raskritie(stroke):
    phrase = str(stroke['phrase'])
    newphrase = phrase
    for key, value in raskrit.items():
        newphrase = newphrase.replace(key, value)  
    return newphrase

def ItFIO(phrase):
    splitted = phrase.split()
    if len(splitted) != 3:
        return False
    res = 0
    for word in splitted:
        parsed_word = morph.parse(word)[0]
        print(morph.parse(word))
        if 'Name' in parsed_word.tag or 'Surn' in parsed_word.tag or 'Patr' in parsed_word.tag:
            res += 1
    return res >= 2

def GetGenderFIO(phrase):
    res = {'femn': 0, 'masc': 0, 'neut': 0}
    for word in phrase.split():
        parsed_word = morph.parse(word)[0]
        if parsed_word.tag.gender:
            res[parsed_word.tag.gender] += 1
    res_sorted = sorted(res, key=lambda x: res[x], reverse=True)
    return res_sorted[0]

def remove_quoted_text(phrase):
    pattern = r'(["«].*?["»])'
    match = re.search(pattern, phrase, re.DOTALL)
    if match:
        return True, match.group(0)
    return False, ''

def ChangeCattle(stroke):
    phrase = str(stroke['phrase'])
    print(phrase)
    ret_quotes = {}
    have, found = remove_quoted_text(phrase)
    while have:
        phrase = phrase.replace(found, f'(*{len(ret_quotes)}*)')
        ret_quotes[len(ret_quotes)] = found
        have, found = remove_quoted_text(phrase)
    if ItFIO(phrase):
        gend = GetGenderFIO(phrase)
        print(gend)
        params = {gend, stroke['filters']['ChangeCattle']}
    else:
        params = {stroke['filters']['ChangeCattle']}
    
    newphrase = ''
    for key, value in matchers.items():
        result = re.search(key, phrase, re.DOTALL)
        if result:
            phrase1 = result.group(0)
            try:
                splitted_phrase = re.split(value, phrase1)[1]
                phrase = phrase.replace(splitted_phrase, f'(*{len(ret_quotes)}*)')
                ret_quotes[len(ret_quotes)] = splitted_phrase
            except IndexError:
                splitted_phrase = ''
        else:
            continue
    for word in phrase.split():
        try:
            parsed_word = morph.parse(word)[0]
            newphrase += restore_capitalization(parsed_word.inflect(params).word, word)+ ' '
        except:
            newphrase += word+ ' '
            continue
    for key, value in ret_quotes.items():
        newphrase = newphrase.replace(f'(*{key}*)', value)
    return newphrase.strip()

def UpperCase(stroke):
    response = HttpResponse()
    phrase = str(stroke['phrase'])
    newphrase = phrase
    if stroke['filters']['ChangeCase'] == 'Upper':
        newphrase = phrase.upper()
    elif stroke['filters']['ChangeCase'] == 'Lower':
        newphrase = phrase.lower()
    return newphrase
    
def connect_objects_to_document(request, pk):
    if request.method == 'POST':
        print(request.POST)
        # Получить список выбранных объектов из запроса
        selected_objects = request.POST.getlist('selectedObjects[]')
        
        print(selected_objects)

        # Получить документ по идентификатору
        document = DocumentsPattern.objects.get(pk=pk)

        # Подключить выбранные объекты к документу
        for obj_id in selected_objects:
            obj = Object.objects.get(id=obj_id)
            doc_obj = DocumentPattern_Objects.objects.filter(document=document, object=obj).first()
            if doc_obj is None:
                doc_obj = DocumentPattern_Objects()
            doc_obj.document = document
            doc_obj.object = obj
            doc_obj.save()

        # Вернуть успешный ответ
        return HttpResponse()
    
def delete_object_from_document(request, pk):
    if request.method == 'POST':
        try:
            print(request.POST)
            # Получить список выбранных объектов из запроса
            object_id = request.POST['object_id']
            
            print(object_id)

            # Получить документ по идентификатору
            document = DocumentsPattern.objects.get(pk=pk)

            # Подключить выбранные объекты к документу
            doc_obj = DocumentPattern_Objects.objects.filter(document=document, object=object_id)
            doc_obj.delete()

            # Вернуть успешный ответ
            return HttpResponse()
        except Exception as e:
            print(e)
            return HttpResponseNotModified('Не удалено. Ошибка удаления. ' + e)
        
     