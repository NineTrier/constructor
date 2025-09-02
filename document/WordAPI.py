from docx import Document
from docx.oxml import OxmlElement, ns

# Пока что не знаю куда разместить эти функции, поэтому лежат просто в файле
# Функция создаёт элемент xml
def create_element(name: str):
    return OxmlElement(name)


# Функция создаёт аттрибут для элемента xml
def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


# Класс API для работы с документами Microsoft Word в xml структуре
class WordAPI:

    def __init__(self, document=None):
        if document is None:
            self.doc = Document()
        else:
            self.doc = Document(document)

    def openDoc(self, doc):
        self.doc = Document(doc)

    def saveDoc(self, path):
        self.doc.save(path)

    def delete_element(self, elem):
        try:
            run = self.doc.paragraphs[0].add_run()
            for bad in run._r.xpath(elem):
                bad.getparent().remove(bad)
        except Exception as exc:
            print(exc, "delete_element")

    def delete_attribute(self, elem, attribute):
        try:
            elem.attrib.pop(attribute)
        except Exception as exc:
            print(exc, "delete_attribute")

    def get_elements_by_tag(self, tag):
        elements = []
        try:
            run = self.doc.add_paragraph().add_run()
            for bad in run._r.xpath(f"//{tag}"):
                elements.append(bad)
            return elements
        except Exception as exc:
            print(exc, "get_elements_by_tag")

    def create_new_doc(self, lxmlBody):
        try:
            body = self.get_elements_by_tag("w:body")[0]
            print(lxmlBody)
            for bad in body:
                body.remove(bad)
            for par in lxmlBody:
                if par.tag != 'p':
                    continue
                print(par)
                res, el = par.to_lxml()
                if(res):
                    body.append(el)
        except Exception as exc:
            print(exc, "create_new_doc")

