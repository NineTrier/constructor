# Класс Документ, представленный xml структурой
import json
import re
from django.conf import settings
from docx.shared import Inches, Mm
from .PropertyParser import *

# Специальные xml вставки от Microsoft, нужны, чтобы сравнивать теги
vst = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
vst1 = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
vst2 = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"

vsts = { 
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}": "w", 
    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}": "r", 
    "{http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas}": "wpc", 
    "{http://schemas.microsoft.com/office/drawing/2014/chartex}": "cx", 
    "{http://schemas.microsoft.com/office/drawing/2015/9/8/chartex}": "cx1", 
    "{http://schemas.microsoft.com/office/drawing/2015/10/21/chartex}": "cx2", 
    "{http://schemas.microsoft.com/office/drawing/2016/5/9/chartex}": "cx3", 
    "{http://schemas.microsoft.com/office/drawing/2016/5/10/chartex}": "cx4", 
    "{http://schemas.microsoft.com/office/drawing/2016/5/11/chartex}": "cx5", 
    "{http://schemas.microsoft.com/office/drawing/2016/5/12/chartex}": "cx6", 
    "{http://schemas.microsoft.com/office/drawing/2016/5/13/chartex}": "cx7", 
    "{http://schemas.microsoft.com/office/drawing/2016/5/14/chartex}": "cx8", 
    "{http://schemas.openxmlformats.org/markup-compatibility/2006}": "mc", 
    "{http://schemas.microsoft.com/office/drawing/2016/ink}": "aink", 
    "{http://schemas.microsoft.com/office/drawing/2017/model3d}": "am3d", 
    "{urn:schemas-microsoft-com:office:office}": "o", 
    "{http://schemas.microsoft.com/office/2019/extlst}": "oel", 
    "{http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing}": "wp14", 
    "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}": "wp", 
    "{urn:schemas-microsoft-com:vml}": "v", 
    "{http://schemas.microsoft.com/office/word/2010/wordml}": "w14", 
    "{http://schemas.microsoft.com/office/word/2012/wordml}": "w15", 
    "{http://schemas.microsoft.com/office/word/2018/wordml/cex}": "w16cex", 
    "{http://schemas.microsoft.com/office/word/2016/wordml/cid}": "w16cid", 
    "{http://schemas.microsoft.com/office/word/2018/wordml}": "w16", 
    "{http://schemas.microsoft.com/office/word/2023/wordml/word16du}": "w16du", 
    "{http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash}": "w16sdtdh", 
    "{http://schemas.microsoft.com/office/word/2015/wordml/symex}": "w16se", 
    "{http://schemas.microsoft.com/office/word/2010/wordprocessingGroup}": "wpg", 
    "{http://schemas.microsoft.com/office/word/2010/wordprocessingInk}": "wpi", 
    "{http://schemas.microsoft.com/office/word/2006/wordml}": "wne", 
    "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}": "wps" 
}
i_id = 0


# Абстрактный класс Элемент
class Element(ABC):

    id: str
    childs: list
    tag: str
    attrib: dict

    # Метод, который выводит текст, содержащийся в элементе
    @abstractmethod
    def text(self):
        pass

    # Метод, который переводит элемент в lxml word для формирования документа Microsoft Word
    @abstractmethod
    def to_lxml(self):
        pass

    @abstractmethod
    def from_json(self, json: dict):
        pass


# Объект, который просто копирует xml структуру элемента, для которого не описан класс
class Else(Element):

    def __init__(self, elem):
        self.tag = elem.tag.split('}')[1]
        self.elem = elem
        self.childs = []
        self.attrib = {}
        self.id = ""

    def text(self):
        return ""

    def to_lxml(self):
        return self.elem

    def from_json(self, json: dict):
        pass


# Класс-паттерн фабрика для создания объектов класса Element
class ElementFactory:

    def __init__(self):
        pass

    def initialize(self, elem):
        tag = elem.tag.split('}')[1]
        if tag.find("Pr") != -1:
            return Properties(elem)
        elif tag == "r":
            return Run(elem)
        elif tag == "hyperlink":
            return Hyperlink(elem)
        elif tag == "p":
            return Paragraph(elem)
        elif tag == "t":
            return Text(elem)
        elif tag == "tbl":
            return Table(elem)
        elif tag == "tr":
            return TableRow(elem)
        elif tag == "tc":
            return TableCol(elem)
        else:
            return Else(elem)

    def initialize_from_json(self, json1: dict):
        if type(json1) == str:
            json_stroke = json.loads(json1)
        else:
            json_stroke = json1
        tag = json_stroke['id'].split("_")[0]
        if tag == "r":
            r = Run(None)
            r.from_json(json_stroke)
            return r
        elif tag == "hypers":
            hyp = Hyperlink(None)
            hyp.from_json(json_stroke)
            return hyp
        elif tag == "p":
            p = Paragraph(None)
            p.from_json(json_stroke)
            return p
        elif tag == "tbl":
            tbl = Table(None)
            tbl.from_json(json_stroke)
            return tbl
        elif tag == "tr":
            tr = TableRow(None)
            tr.from_json(json_stroke)
            return tr
        elif tag == "tc":
            tc = TableCol(None)
            tc.from_json(json_stroke)
            return tc
        else:
            return ""

    # Функция возвращает потомков Элемента elem.
    # Пока что эта функция в классе фабрика, что не правильно, но не знаю куда лучше разместить
    def get_childs(self, elem):
        global i_id
        childs = []
        for son in elem:
            childs.append(self.initialize(son))
            childs[-1].id = f"{son.tag.split('}')[1]}_{i_id}"
            i_id += 1
        return childs


# Элемент свойства
class Properties(Element):

    def __init__(self, elem):
        self.id = ""
        if elem is None:
            self.tag = ""
            self.pr = [PropertiesCollection(None)]
            self.attrib = {}
            self.childs = []
        else:
            self.tag = elem.tag.split('}')[1]
            self.pr = [PropertiesCollection(elem, self.tag)]
            self.childs = []
            self.attrib = {f"{vsts[i.split('}')[0]+'}']}:" + i.split('}')[1]: elem.attrib[i] for i in elem.attrib}
            # Заполняем свойства
            for prop in elem:
                # Если свойства древовидные, то они тоже будут элементами класса Properties
                if prop.tag == f"{vst}rPr" or prop.tag == f"{vst}sectPr":
                    self.pr.append(Properties(prop))

    # Перевод в lxml
    def to_lxml(self):
        Pr = create_element(f"w:{self.tag}")
        for k in list(self.attrib.keys()):
            create_attribute(Pr, f"{k}", self.attrib[k])
        for child in self.pr:
            if type(child) == Properties:
                ch = child.to_lxml()
                Pr.append(ch)
            elif type(child) == PropertiesCollection:
                for ch in child.to_lxml():
                    Pr.append(ch)
        return Pr

    def text(self):
        return ""

    # Перевод свойств в css
    def to_css(self):
        css = """""" # Строка css сначала пустая
        for ch in self.pr:
            css += ch.to_css()
        return css

    # Получение элемента из JSON объекта, нужен когда закончили редактирование в html странице
    def from_json(self, json_stroke: dict):
        propCol = PropertiesCollection(None)
        propCol.from_json(json_stroke)
        self.pr = [propCol]

    def __eq__(self, other):
        if self.pr == other.pr:
            return True
        else:
            return False


# Класс Ссылка
class Hyperlink(Element):
    def __init__(self, elem):
        self.id = ""
        self.tag = 'hyperlink'
        if elem is None:
            self.pr = Properties(None)
            self.attrib = {}
            self.childs = []
        else:
            self.attrib = {f"{vsts[i.split('}')[0]+'}']}:" + i.split('}')[1]: elem.attrib[i] for i in elem.attrib}
            self.childs = ElementFactory().get_childs(elem)
            self.pr = Properties(elem[0][0])

    def text(self):
        text = ""
        for child in self.childs:
            text += child.text()
        return text

    def to_lxml(self):
        hyperlink = create_element("w:hyperlink")
        for k in list(self.attrib.keys()):
            create_attribute(hyperlink, f"{k}", self.attrib[k])
        for child in self.childs:
            try:
                hyperlink.append(child.to_lxml())
            except Exception as exc:
                print(exc, "Hyperlink to lxml")
        return hyperlink

    def from_json(self, json1: dict):
        if type(json1) == str:
            json_stroke = json.loads(json1)
        else:
            json_stroke = json1
        self.childs = []
        r = Run(None)
        r.from_json(json_stroke)
        r.id = f'r_{len(self.childs)}'
        self.childs.append(r)
        idd = json_stroke['attrib'].split('/')[0]
        hist = json_stroke['attrib'].split('/')[1]
        self.attrib = {'r:id': idd, 'w:history': hist}


# Класс Текстовый элемент
class Text(Element):
    def __init__(self, elem):
        self.id = ""
        self.tag = 't'
        if elem is None:
            self.textElem = ""
            self.attrib = {}
        else:
            self.textElem = elem.text
            self.attrib = {i.split('}')[1]: elem.attrib[i] for i in elem.attrib}

    def text(self):
        return self.textElem

    def to_lxml(self):
        try:
            t = create_element("w:t")
            for k in list(self.attrib.keys()):
                create_attribute(t, f"xml:{k}", self.attrib[k])
            t.text = self.textElem
            return t
        except Exception as exc:
            print(exc, "Text to lxml", self.textElem)

    def from_json(self, json1: dict):
        if type(json1) == str:
            json_stroke = json.loads(json1)
        else:
            json_stroke = json1
        self.textElem = json_stroke['text']
        self.attrib['space'] = 'preserve'


# Класс пробег по параграфу
class Run(Element):

    def __init__(self, elem):
        self.id = ""
        self.tag = 'r'
        self.hidden = False
        if elem is None:
            self.pr = Properties(None)
            self.attrib = {}
            self.childs = []
        else:
            self.attrib = {f"{vsts[i.split('}')[0]+'}']}:" + i.split('}')[1]: elem.attrib[i] for i in elem.attrib}
            self.childs = ElementFactory().get_childs(elem)
            self.pr = Properties(elem[0]) if len(elem) > 0 and elem[0].tag.split('}')[1] == "rPr" else ""

    def text(self):
        t = ""
        for son in self.childs:
            if re.fullmatch("t_\d*", son.id):
                t += son.text()
        return t

    def to_lxml(self):
        if self.hidden:
            return False, None
        r = create_element("w:r")
        for k in list(self.attrib.keys()):
            create_attribute(r, f"{k}", self.attrib[k])
        for child in self.childs:
            try:
                r.append(child.to_lxml())
            except Exception as exc:
                print(exc, "Run to lxml", child)
        return True, r

    def from_json(self, json1: dict):
        if type(json1) == str:
            json_stroke = json.loads(json1)
        else:
            json_stroke = json1
        self.childs = []
        self.hidden = json_stroke['hidden']
        rPr = Properties(None)
        rPr.tag = 'rPr'
        rPr.id = f'rPr_{len(self.childs)}'
        rPr.from_json(json_stroke['style'])
        self.childs.append(rPr)
        text = Text(None)
        text.from_json(json_stroke)
        self.childs.append(text)


# Класс параграф
class Paragraph(Element):

    def __init__(self, elem):
        self.id = ""
        self.tag = 'p'
        self.hidden = False
        if elem is None:
            self.pr = Properties(None)
            self.attrib = {}
            self.childs = []
        else:
            self.pr = Properties(elem[0]) if len(elem) > 0 and elem[0].tag.split('}')[1] == "pPr" else ""
            self.attrib = {f"{vsts[i.split('}')[0]+'}']}:" + i.split('}')[1]: elem.attrib[i] for i in elem.attrib}
            self.childs = ElementFactory().get_childs(elem)

    def text(self):
        text = ""
        for child in self.childs:
            text += child.text()
        return text

    def to_lxml(self):
        if self.hidden:
            return False, None
        p = create_element("w:p")
        for k in list(self.attrib.keys()):
            create_attribute(p, f"{k}", self.attrib[k])
        for child in self.childs:
            try:
                if child.tag == 'r':
                    yes, element = child.to_lxml()
                    if yes:
                        p.append(element)
                else:
                    p.append(child.to_lxml())
            except Exception as exc:
                print(exc, "Paragraph to lxml", child)
        return True, p

    def from_json(self, json1):
        if type(json1) == str:
            json_stroke = json.loads(json1)
        else:
            json_stroke = json1
        self.hidden = json_stroke['hidden']
        self.childs = []
        pPr = Properties(None)
        pPr.tag = 'pPr'
        pPr.id = f'pPr_{len(self.childs)}'
        pPr.from_json(json_stroke['style'])
        self.childs.append(pPr)
        for child in json_stroke['childs']:
            json_childs = json.loads(child)
            if json_childs['class'] == 'runs':
                run = Run(None)
                run.from_json(json_childs)
                self.childs.append(run)
            elif json_childs['class'] == 'hypers':
                hyper = Hyperlink(None)
                hyper.from_json(child)
                self.childs.append(json_childs)
            elif json_childs['class'] == 'img':
                print('Это картинка')


# Класс таблица
class Table(Element):

    def __init__(self, elem):
        self.id = ""
        self.tag = "tbl"
        if elem is None:
            self.pr = Properties(None)
            self.attrib = {}
            self.childs = []
            self.tableGrid = []
        else:
            self.pr = Properties(elem[0]) if len(elem) > 0 and elem[0].tag.split('}')[1] == "tblPr" else ""
            self.attrib = {f"{vsts[i.split('}')[0]+'}']}:{i.split('}')[1]}": elem.attrib[i] for i in elem.attrib}
            self.childs = ElementFactory().get_childs(elem)
            self.tableGrid = [{f"{vsts[i.split('}')[0]+'}']}:{i.split('}')[1]}": gridCol.attrib[i] for i in gridCol.attrib} for gridCol in elem[1]]

    def text(self):
        text = ""
        for child in self.childs:
            text += child.text()
        return text

    def to_lxml(self):
        tbl = create_element("w:tbl")
        for k in list(self.attrib.keys()):
            create_attribute(tbl, f"{k}", self.attrib[k])
        for child in self.childs:
            try:
                tbl.append(child.to_lxml())
                if child.tag.find("Pr") != -1:
                    tableGrid = create_element("w:tblGrid")
                    for gridCol in self.tableGrid:
                        for k, v in gridCol.items():
                            gridCol = create_element("w:gridCol")
                            create_attribute(gridCol, "w:w", v)
                            tableGrid.append(gridCol)
                    tbl.append(tableGrid)
            except Exception as exc:
                print(exc, "Table to lxml", child)
        return tbl

    def from_json(self, json1: dict):
        if type(json1) == str:
            json_stroke = json.loads(json1)
        else:
            json_stroke = json1
        self.childs = []
        tblPr = Properties(None)
        tblPr.tag = 'tblPr'
        tblPr.id = f'tblPr_{len(self.childs)}'
        tblPr.from_json(json_stroke['style'])
        self.childs.append(tblPr)
        for child_str in json_stroke['childs']:
            child = json.loads(child_str)
            if child['id'].split('_')[0] == 'tr':
                row = TableRow(None)
                row.from_json(child)
                self.childs.append(row)
        self.tableGrid = [{"w:w": i.split(':')[-1]} for i in json_stroke['tableGrid'].split('|')]


# Класс строка таблицы
class TableRow(Element):

    def __init__(self, elem):
        self.id = ""
        self.tag = "tr"
        if elem is None:
            self.pr = Properties(None)
            self.attrib = {}
            self.childs = []
        else:
            self.pr = Properties(elem[0]) if len(elem) > 0 and elem[0].tag.split('}')[1] == "tblPr" else ""
            self.attrib = {f"{vsts[i.split('}')[0]+'}']}:" + i.split('}')[1]: elem.attrib[i] for i in elem.attrib}
            self.childs = ElementFactory().get_childs(elem)

    def text(self):
        text = ""
        for child in self.childs:
            text += child.text()
        return text

    def to_lxml(self):
        p = create_element("w:tr")
        for k in list(self.attrib.keys()):
            create_attribute(p, f"{k}", self.attrib[k])
        for child in self.childs:
            try:
                p.append(child.to_lxml())
            except Exception as exc:
                print(exc, "TableRow to lxml", child)
        return p

    def from_json(self, json1: dict):
        if type(json1) == str:
            json_stroke = json.loads(json1)
        else:
            json_stroke = json1
        self.childs = []
        trPr = Properties(None)
        trPr.tag = 'trPr'
        trPr.id = f'trPr_{len(self.childs)}'
        trPr.from_json(json_stroke['style'])
        self.childs.append(trPr)
        for child_str in json_stroke['childs']:
            child = json.loads(child_str)
            if child['id'].split('_')[0] == 'tc':
                col = TableCol(None)
                col.from_json(child)
                self.childs.append(col)


# Класс ячейка таблицы (Но почему-то назвал как колонка таблицы, но по логике это ячейка)
class TableCol(Element):

    def __init__(self, elem):
        self.id = ""
        self.tag = "tc"
        if elem is None:
            self.pr = Properties(None)
            self.attrib = {}
            self.childs = []
        else:
            self.pr = Properties(elem[0]) if len(elem) > 0 and elem[0].tag.split('}')[1] == "tcPr" else Properties(None)
            self.attrib = {f"{vsts[i.split('}')[0]+'}']}:" + i.split('}')[1]: elem.attrib[i] for i in elem.attrib}
            self.childs = ElementFactory().get_childs(elem)

    def text(self):
        text = ""
        for child in self.childs:
            text += child.text()
        return text

    def to_lxml(self):
        p = create_element("w:tc")
        for k in list(self.attrib.keys()):
            create_attribute(p, f"{k}", self.attrib[k])
        for child in self.childs:
            try:
                p.append(child.to_lxml())
            except Exception as exc:
                print(exc, "TableCol to lxml", child.pr)
        return p

    def from_json(self, json1: dict):
        if type(json1) == str:
            json_stroke = json.loads(json1)
        else:
            json_stroke = json1
        self.childs = []
        tcPr = Properties(None)
        tcPr.tag = 'tcPr'
        tcPr.id = f'tcPr_{len(self.childs)}'
        tcPr.from_json(json_stroke['style'])
        self.childs.append(tcPr)
        par = Paragraph(None)
        par.from_json(json_stroke['paragraph'])
        self.childs.append(par)


# Класс документ
class Document:

    def __init__(self, path=None):
        global i_id
        i_id = 0
        self.images = {}
        if path is None:
            self.wa = WordAPI()
            self.childs = []
        else:
            self.wa = WordAPI(path)
            try:
                body = self.wa.get_elements_by_tag("w:body")[0]
            except TypeError:
                self.childs = []
                return
            self.childs = ElementFactory().get_childs(body)

    def save(self, path):
        self.wa.create_new_doc(self.childs)
        self.wa.saveDoc(path)
        for key, value in self.images.items():
            self.wa.doc.paragraphs[int(key)]._element.clear_content()
            print(value['src'])
            print(settings.MEDIA_ROOT)
            self.wa.doc.paragraphs[int(key)].add_run().add_picture(f"{settings.MEDIA_ROOT}/{value['src'].split('media')[1]}", width=Mm(int(value['width'])), height=Mm(int(value['height'])))
        self.wa.saveDoc(path)


    def from_json(self, json_stroke: dict):
        self.childs = []
        self.images = json_stroke['images']
        print(self.images)
        for elem in json_stroke['elements']:
            self.childs.append(ElementFactory().initialize_from_json(elem))
        sectPr = Properties(None)
        sectPr.from_json({i: i for i in json_stroke['sectPr'].split('|')})
        self.childs.append(sectPr)



