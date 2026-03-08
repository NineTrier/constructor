import json
import base64
import re
import json

from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic import CreateView
from .models import DocumentsPattern, DocumentPattern_Objects,Fonts, SavedElements, VariableBlock, DocType, Document_ParentDocument, Document_VariableBlock
from database_manager.models import Object, ObjectLinkMeta, Object_ParentObject, Parameter
from .forms import DocumentForm
from django.core.files.storage import FileSystemStorage
import os
from django.conf import settings
from django.http import HttpRequest, HttpResponse, Http404, HttpResponseNotModified, HttpResponseForbidden, JsonResponse, HttpResponseBadRequest
from docx import Document
from .Document import Document

from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.contrib.auth.decorators import login_required
from transliterate import translit
from transliterate.decorators import transliterate_function
from user_manager.models import Profile, user_directory_path
from pathlib import Path, PurePosixPath
from typing import Any, Dict, Optional, Tuple

import logging
logger = logging.getLogger(__name__)
from .token_resolver import (
    ExportFinalizeError,
    TokenResolveError,
    TokenResolverService,
    build_doc_token_index,
    finalize_document_json_for_export,
)

import pymorphy3
from pymorphy3.shapes import restore_capitalization

raskrit = {
    'ООО': 'Общество с ограниченной ответственностью',
    'ИП':'Индивидуальный предприниматель',
    'ПАО': 'Публичное акционерное общество',
    'АО': 'Акционерное общество'
}

matchers ={ 
    r'\s?Арбитражн\w*\sсуд\w?': r'\s?Арбитражн\w*\sсуд\w?\s',
    r'\s?Обществ\w*\sс\sограниченной\sответственностью\w*': r'\s?Обществ\w+?\s',
    r'\s?Индивидуальн\w*\sпредпринимател\w?': r'\s?Индивидуальн\w*\sпредпринимател\w?\s',
    r'\s?Акционерн\w*\sобществ\w?': r'\s?Акционерн\w*\sобществ\w?\s',
}

morph = pymorphy3.MorphAnalyzer()


# Helpers for filesystem operations ----------------------------------------
def _user_media_dir(user_id, *parts):
    """Ensure MEDIA_ROOT/documents/user_<id>/... exists and return Path."""
    base = Path(settings.MEDIA_ROOT) / "documents" / f"user_{user_id}"
    if parts:
        base = base.joinpath(*[str(part) for part in parts])
    base.mkdir(parents=True, exist_ok=True)
    return base


def _user_media_rel(user_id, *parts):
    """Return POSIX relative path under user documents folder."""
    return PurePosixPath("documents").joinpath(f"user_{user_id}", *[str(part) for part in parts])


def _copy_document_for_user(request, document_id):
    """
    Internal helper that clones a document template for the current user.

    Returns tuple: (success: bool, document instance or None, error message or None)
    """
    try:
        profile = Profile.for_user(request.user)

        document_to_copy = DocumentsPattern.objects.get(id=document_id)
        document_parent_qs = Document_ParentDocument.objects.filter(parent=document_to_copy, userRequested=profile)
        objects = DocumentPattern_Objects.objects.filter(document=document_to_copy)
        variables = Document_VariableBlock.objects.filter(document=document_to_copy.id)

        if document_parent_qs:
            document = DocumentsPattern.objects.get(id=document_parent_qs[0].document.id)
            document_parent = document_parent_qs[0]
        else:
            document = DocumentsPattern()
            document_parent = Document_ParentDocument()
            document_parent.document = document
            document_parent.parent = document_to_copy
            document_parent.userRequested = profile
            document_parent.parentDocumentChanged = False

        document.name = document_to_copy.name[:40]
        document.owner = profile
        document.type = document_to_copy.type
        document.description = document_to_copy.description
        document.picture = "noimage.jpeg"
        document.json = document_to_copy.json

        latest_document = DocumentsPattern.objects.order_by('-id').first()
        next_id = (latest_document.id if latest_document else 0) + 1
        target_dir = _user_media_dir(request.user.id, next_id)
        file_name = f"{translit_russian(document.name)}{translit_russian(document.owner.__str__())}.docx"
        destination_path = target_dir / file_name

        source_candidates = []
        try:
            source_candidates.append(Path(document_to_copy.file.path))
        except (ValueError, AttributeError):
            pass
        relative_raw = str(document_to_copy.file)
        if relative_raw:
            relative_normalized = relative_raw.replace('\\', '/')
            source_candidates.append(Path(settings.MEDIA_ROOT) / relative_normalized)
            source_candidates.append(Path(settings.BASE_DIR) / relative_normalized)

        source_path = next((p for p in source_candidates if p and p.is_file()), None)
        if not source_path:
            raise FileNotFoundError(
                f"Source document file not found. Tried paths: "
                f"{', '.join(str(p) for p in source_candidates if p)}"
            )

        document_api = Document(str(source_path))
        document_api.save(str(destination_path))
        document.file = _user_media_rel(request.user.id, next_id, file_name).as_posix()
        document.documentOfOrganisation = False
        document.save()
        document_parent.save()

        for document_variable in variables:
            new_document_variable = Document_VariableBlock.objects.filter(
                document=document,
                variable=document_variable.variable
            ).first() or Document_VariableBlock()
            new_document_variable.variable = document_variable.variable
            new_document_variable.document = document
            new_document_variable.save()

        for document_object in objects:
            new_document_object = DocumentPattern_Objects.objects.filter(
                object=document_object.object,
                document=document
            ).first() or DocumentPattern_Objects()
            new_document_object.document = document
            new_document_object.object = document_object.object
            new_document_object.save()

        return True, document, None

    except Exception as exc:
        logger.exception("Failed to copy document %s for user %s", document_id, request.user)
        return False, None, str(exc)


def _api_v1_error(code: str, message: str, *, status: int = 400, details: Optional[Dict[str, Any]] = None) -> JsonResponse:
    return JsonResponse(
        {
            "error": {
                "code": code,
                "message": message,
                "details": details or {},
            }
        },
        status=status,
    )


def _parse_json_body(request) -> Optional[Dict[str, Any]]:
    body = request.body.decode("utf-8").strip() if request.body else ""
    if not body:
        return {}
    try:
        payload = json.loads(body)
    except json.JSONDecodeError:
        return None
    if not isinstance(payload, dict):
        return None
    return payload


def _can_access_document(user, document: DocumentsPattern) -> bool:
    owner_user = getattr(getattr(document, "owner", None), "user", None)
    if owner_user is not None and owner_user == user:
        return True
    return bool(user.is_superuser or user.has_perm("document.change_documentspattern"))


def _build_export_payload_for_docx(*, document: DocumentsPattern) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    source_json = document.json if isinstance(document.json, dict) else {}
    if "elements" not in source_json or "sectPr" not in source_json:
        raise ValueError("Document JSON is not ready for export finalization.")
    resolver = TokenResolverService(logger=logger)
    return finalize_document_json_for_export(
        document_id=int(document.id),
        document_json=source_json,
        context=source_json.get("dbm_context") or source_json.get("context") or {},
        options={
            "maxDepth": 8,
            "aggregationMode": "first",
            "joiner": ", ",
            "validateOnly": False,
            "includeTrace": False,
        },
        resolver=resolver,
    )


# Класс, который помогает создавать новый записи в базу данных Documents
# и открывает страницу с добавлением новых документов
class DocumentCreate(CreateView):
    model = DocumentsPattern
    form_class = DocumentForm

    template_name = 'document/document_create.html'
    
    def get_context_data(self, **kwargs):
        context = super(DocumentCreate, self).get_context_data(**kwargs)
        if self.request.user.is_authenticated:
            context['profile'] = Profile.for_user(self.request.user)
        context['documents'] = DocumentsPattern.objects.all()
        context['title'] = "Загрузка документа"
        return context
    
    def post(self, request, *args, **kwargs):
        document = DocumentsPattern()
        document.name = request.POST['name']
        document.description = request.POST['description']
        document.type = DocType.objects.get(id=request.POST['type'])
        document.owner = Profile.for_user(request.user)
        file = request.FILES.get('file')
        try:
            latest_document = DocumentsPattern.objects.latest('id')
            next_id = latest_document.id + 1
        except DocumentsPattern.DoesNotExist:
            next_id = 1
        user_dir = _user_media_dir(request.user.id, next_id)
        safe_name = translit_russian(file.name)
        fs = FileSystemStorage(location=str(user_dir))
        fs.save(safe_name, file)
        document.file = _user_media_rel(request.user.id, next_id, safe_name).as_posix()
        document.save()
        response = redirect(f"/document/view?id={document.id}&type=1")
        return response

def translit_russian(text):
    """Выполняет транслитерацию русского текста в латиницу для названий файлов"""
    try:
        translited = translit(text, reversed=True)
        translited = str(translited).replace(' ', '_')
        return translited
    except Exception as exc:
        print(exc)
        return text
    

def create_New_Document(request):
    """Обработчик запроса для создания документа"""
    if request.method == 'POST':
        document = DocumentsPattern()
        document.name = request.POST['name']
        document.description = request.POST['description']
        document.type = DocType.objects.get(id=request.POST['type'])
        document.owner = Profile.for_user(request.user)
        document.picture = f"noimage.jpeg"
        print(document)
        try:
            latest_document = DocumentsPattern.objects.latest('id')
            next_id = latest_document.id + 1
        except DocumentsPattern.DoesNotExist:
            next_id = 1
        target_dir = _user_media_dir(request.user.id, next_id)
        file_name = f"{translit_russian(document.name)}{translit_russian(document.owner.__str__())}"[:50]
        destination_path = target_dir / f"{file_name}.docx"
        document.file = _user_media_rel(request.user.id, next_id, f"{file_name}.docx").as_posix()
        blank_template = Path(settings.BASE_DIR).parent / 'blank.docx'
        doc_file = Document(str(blank_template))
        doc_file.save(str(destination_path))
        document.save()
        return redirect(f"/document/view?id={document.id}&type=1")
    form = DocumentForm()
    context = {
        'form': form,
        'title': 'Создание документа'
    }
    if request.user.is_authenticated:
        context['profile'] = Profile.for_user(request.user)
    return render(request, 'document/new_document.html', context)

@csrf_exempt
def SaveCover(request):
    """Обработчик запроса для загрузки обложки"""
    try:
        id = request.POST['id']
        document = DocumentsPattern.objects.get(id=id)
        img = request.POST['img']
        img = str(img).replace('data:image/png;base64,', '')
        img = str(img).replace(' ', '+')
        dat = base64.decodebytes(img.encode('utf-8'))
        target_dir = _user_media_dir(request.user.id, document.id)
        file_name = f"{translit_russian(document.name)}Cover.png"
        file_path = target_dir / file_name
        file_path.write_bytes(dat)
        document.picture = _user_media_rel(request.user.id, document.id, file_name).as_posix()
        document.save() 
        return HttpResponse()
    except Exception as exc:
        print(exc)
        return HttpResponseNotModified()

@csrf_exempt
def SaveImage(request):
    """Обработчик запроса для загрузки картинок в документ"""
    try:
        id = request.POST['id']
        document = DocumentsPattern.objects.get(id=id)
        img = request.FILES['img']
        target_dir = _user_media_dir(request.user.id, document.id)
        image_name = f"{request.POST['idImage']}.png"
        file_path = target_dir / image_name
        if file_path.exists():
            file_path.unlink()
        storage_path = _user_media_rel(request.user.id, document.id, image_name).as_posix()
        fs = FileSystemStorage()
        fs.save(storage_path, img)
        response = HttpResponse()
        response['img'] = f"{settings.MEDIA_URL}{storage_path}"
        return response
    except Exception as exc:
        print(exc)
        return HttpResponseNotModified()

# Функция, которая позволяет скачать файл, загруженный на сервер
def download(request):
    """Обработчик запроса для выгрузки документа"""
    document = DocumentsPattern.objects.filter(id=request.GET.get('id'))[0]
    if not _can_access_document(request.user, document):
        return HttpResponseForbidden()
    document.downloadsTimes = int(document.downloadsTimes) + 1
    document.save()
    filepath = document.file
    file_path = os.path.join(settings.MEDIA_ROOT, str(filepath))
    source_json = document.json if isinstance(document.json, dict) else {}
    if "elements" in source_json and "sectPr" in source_json:
        try:
            export_payload, resolve_result = _build_export_payload_for_docx(document=document)
            document_export = Document()
            document_export.from_json(export_payload)
            document_export.save(file_path)
            summary = resolve_result.get("summary", {})
            logger.info(
                "docx_export_resolved document_id=%s tokens=%s ok=%s errors=%s warnings=%s",
                document.id,
                summary.get("tokens_total", 0),
                summary.get("ok", 0),
                summary.get("errors", 0),
                summary.get("warnings", 0),
            )
        except ExportFinalizeError as exc:
            logger.warning("DOCX export blocked for document %s: %s", document.id, str(exc))
            return JsonResponse(
                {
                    "error": {
                        "code": "TOKEN_RESOLVE_FAILED",
                        "message": str(exc),
                        "details": {
                            "results": exc.results,
                        },
                    }
                },
                status=400,
            )
        except Exception as exc:
            logger.exception("DOCX export failed for document %s", document.id)
            return JsonResponse(
                {
                    "error": {
                        "code": "SERVER_ERROR",
                        "message": "Не удалось подготовить DOCX для выгрузки.",
                        "details": {"exc": str(exc)},
                    }
                },
                status=500,
            )
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
            response['Content-Disposition'] = 'inline; filename=' + f'{translit_russian(document.name)}.docx'
            return response
    raise Http404

@csrf_exempt
def DeleteSavedElement(request):
    """Обработчик запроса для удаления сохраненного элемента"""
    response = redirect('/')
    request_data = request.body
    stroke = json.loads(request_data)
    if stroke['id'] == '-1':
        return response
    SavedElements.objects.filter(id=stroke['id']).delete()
    return response

@csrf_exempt
def DeleteVariable(request):
    """Обработчик запроса для удаления переменной"""
    response = HttpResponse()
    request_data = request.body
    stroke = json.loads(request_data)
    if stroke['id'] == '-1':
        return response
    VariableBlock.objects.filter(id=stroke['id']).delete()
    return response

@csrf_exempt
def CreateDocType(request):
    """Обработчик запроса для создания типа документа"""
    try:
        response = HttpResponse()
        request_data = request.body
        stroke = json.loads(request_data)
        type = DocType()
        type.name = stroke['name']
        type.save()
        response['id'] = type.id
        return response
    except Exception as exc:
        print(exc)
        return HttpResponseNotModified()

@csrf_exempt
def DeleteDocType(request):
    """Обработчик запроса для удаления типа документа"""
    try:
        response = HttpResponse()
        request_data = request.body
        stroke = json.loads(request_data)
        type = DocType.objects.get(id=stroke['id'])
        type.delete()
        return response
    except Exception as exc:
        print(exc)
        return HttpResponseNotModified()

@csrf_exempt
def SaveVariable(request):
    """Обработчик запроса для сохранения переменной в базу данных"""
    request_data = request.body
    stroke = json.loads(request_data)
    doc_variables = Document_VariableBlock.objects.filter(document=stroke['id_doc'])
    if doc_variables == None:
        doc_variables = []
    for doc_variable in doc_variables:
        if doc_variable.variable.name == stroke['name']:
            response = HttpResponseNotModified()
            response['MessageOfError'] = 'Переменная с таким именем существует.'.encode('utf-8')
            response['TypeOfError'] = 'Ошибка переменных'.encode('utf-8')
            return response
    response = HttpResponse()
    if stroke['id'] == '-1':
        variable = VariableBlock()
    else:
        variable = VariableBlock.objects.filter(id=stroke['id'])[0]
    variable.name = stroke['name']
    variable.meaning = stroke['value']
    variable.save()
    new_doc_variable = Document_VariableBlock()
    new_doc_variable.document = DocumentsPattern.objects.filter(id=stroke['id_doc'])[0]
    new_doc_variable.variable = variable
    new_doc_variable.save()
    response['id'] = variable.id
    return response

@csrf_exempt
def CopyDocument(request, id=None):
    """Обработчик запроса для копирования документа пользователю"""
    try:
        response = HttpResponse()
        request_data = request.body
        if id == None:
            stroke = json.loads(request_data)
        else:
            stroke = {'id': id}
        print(stroke)
        profile = Profile.for_user(request.user)
        documentToCopy = DocumentsPattern.objects.get(id=stroke['id'])
        objects = DocumentPattern_Objects.objects.filter(document=documentToCopy)
        variables = Document_VariableBlock.objects.filter(document=documentToCopy.id)
        document = DocumentsPattern()
        document.name = documentToCopy.name + '-копия'
        document.owner = profile
        document.type = documentToCopy.type
        document.description = documentToCopy.description
        document.picture = f"noimage.jpeg"
        document.json = documentToCopy.json
        latest_document = DocumentsPattern.objects.order_by('-id').first()
        next_id = (latest_document.id if latest_document else 0) + 1
        target_dir = _user_media_dir(request.user.id, next_id)
        file_name = f"{translit_russian(document.name)}{translit_russian(document.owner.__str__())}.docx"
        destination_path = target_dir / file_name
        source_path = getattr(documentToCopy.file, "path", None)
        doc_file = Document(source_path)
        doc_file.save(str(destination_path))
        document.file = _user_media_rel(request.user.id, next_id, file_name).as_posix()
        document.documentOfOrganisation = False
        if document.save():
            print('Документ добавлен')
            response['id'] = document.id
            for document_variable in variables:
                print(document_variable)
                new_document_variable = Document_VariableBlock.objects.filter(document=document, variable=document_variable.variable).first()
                if not new_document_variable:
                    new_document_variable = Document_VariableBlock()
                new_document_variable = Document_VariableBlock()
                new_document_variable.variable = document_variable.variable
                new_document_variable.document = document
                new_document_variable.save()
            for document_object in objects:
                new_document_object = DocumentPattern_Objects.objects.filter(object=document_object.object, document=document).first()
                if not new_document_object:
                    new_document_object = DocumentPattern_Objects()
                new_document_object.document = document
                new_document_object.object = document_object.object
                new_document_object.save()
            return response
        else:
            print('документ не сохранен')
            return HttpResponseNotModified()
    except Exception as exc:
        print(f"Не получилось {exc}")
        return HttpResponseNotModified()

@login_required
@require_POST
def toggle_document_organisation(request):
    """Toggle the organisation flag for a document."""
    try:
        try:
            payload = json.loads(request.body.decode('utf-8'))
        except (AttributeError, json.JSONDecodeError):
            payload = request.POST
        doc_id = payload.get('id')
        if not doc_id:
            message = "Document id is required."
            logger.warning("toggle_document_organisation rejected: %s (user=%s)", message, request.user)
            return HttpResponseBadRequest(message)
        document = get_object_or_404(DocumentsPattern, pk=int(doc_id))
        if document.owner.user != request.user and not request.user.has_perm('document.toggle_document_organisation'):
            return HttpResponseForbidden()
        document.documentOfOrganisation = not bool(document.documentOfOrganisation)
        document.save(update_fields=['documentOfOrganisation', 'lastUpdate'])
        return JsonResponse({'documentOfOrganisation': document.documentOfOrganisation})
    except Exception as exc:
        print(exc)
        return HttpResponseNotModified()


@login_required
@require_POST
def api_v1_resolve_tokens(request):
    payload = _parse_json_body(request)
    if payload is None:
        return _api_v1_error("VALIDATION_ERROR", "Ожидался JSON payload.", details={"field": "body"})

    document_id_raw = payload.get("document_id")
    if document_id_raw in (None, ""):
        return _api_v1_error("VALIDATION_ERROR", "Не указан document_id.", details={"field": "document_id"})
    try:
        document_id = int(document_id_raw)
    except (TypeError, ValueError):
        return _api_v1_error("VALIDATION_ERROR", "Некорректный document_id.", details={"field": "document_id"})

    document = DocumentsPattern.objects.filter(id=document_id).first()
    if document is None:
        return _api_v1_error("NOT_FOUND", "Документ не найден.", status=404)
    if not _can_access_document(request.user, document):
        return _api_v1_error("PERMISSION_DENIED", "Недостаточно прав для доступа к документу.", status=403)

    tokens = payload.get("tokens", [])
    if not isinstance(tokens, list):
        return _api_v1_error("VALIDATION_ERROR", "Поле tokens должно быть массивом.", details={"field": "tokens"})
    tokens_list = [str(item or "").strip() for item in tokens if str(item or "").strip()]
    if len(tokens_list) > 500:
        return _api_v1_error(
            "VALIDATION_ERROR",
            "Превышен лимит токенов в одном запросе (максимум 500).",
            details={"tokens_limit": 500},
        )

    context = payload.get("context", {})
    if context is None:
        context = {}
    if not isinstance(context, dict):
        return _api_v1_error("VALIDATION_ERROR", "Поле context должно быть объектом.", details={"field": "context"})

    options = payload.get("options", {})
    if options is None:
        options = {}
    if not isinstance(options, dict):
        return _api_v1_error("VALIDATION_ERROR", "Поле options должно быть объектом.", details={"field": "options"})

    include_trace = bool(options.get("includeTrace", False))
    resolver = TokenResolverService(logger=logger)
    try:
        result = resolver.resolve_tokens(
            document_id=document_id,
            context=context,
            tokens=tokens_list,
            options=options,
            include_trace=include_trace,
        )
    except TokenResolveError as exc:
        return _api_v1_error(exc.code, exc.message, details=exc.details)
    except Exception as exc:
        logger.exception("resolve_tokens API failed for document %s", document_id)
        return _api_v1_error(
            "SERVER_ERROR",
            "Не удалось выполнить резолв токенов.",
            status=500,
            details={"exc": str(exc)},
        )

    return JsonResponse(result)


@login_required
@require_POST
def api_v1_prefetch_graph(request):
    payload = _parse_json_body(request)
    if payload is None:
        return _api_v1_error("VALIDATION_ERROR", "Ожидался JSON payload.", details={"field": "body"})

    document_id_raw = payload.get("document_id")
    if document_id_raw in (None, ""):
        return _api_v1_error("VALIDATION_ERROR", "Не указан document_id.", details={"field": "document_id"})
    try:
        document_id = int(document_id_raw)
    except (TypeError, ValueError):
        return _api_v1_error("VALIDATION_ERROR", "Некорректный document_id.", details={"field": "document_id"})

    document = DocumentsPattern.objects.filter(id=document_id).first()
    if document is None:
        return _api_v1_error("NOT_FOUND", "Документ не найден.", status=404)
    if not _can_access_document(request.user, document):
        return _api_v1_error("PERMISSION_DENIED", "Недостаточно прав для доступа к документу.", status=403)

    tokens = payload.get("tokens", [])
    if not isinstance(tokens, list):
        return _api_v1_error("VALIDATION_ERROR", "Поле tokens должно быть массивом.", details={"field": "tokens"})
    tokens_list = [str(item or "").strip() for item in tokens if str(item or "").strip()]
    if len(tokens_list) > 500:
        return _api_v1_error(
            "VALIDATION_ERROR",
            "Превышен лимит токенов в одном запросе (максимум 500).",
            details={"tokens_limit": 500},
        )

    context = payload.get("context", {})
    if context is None:
        context = {}
    if not isinstance(context, dict):
        return _api_v1_error("VALIDATION_ERROR", "Поле context должно быть объектом.", details={"field": "context"})

    options = payload.get("options", {})
    if options is None:
        options = {}
    if not isinstance(options, dict):
        return _api_v1_error("VALIDATION_ERROR", "Поле options должно быть объектом.", details={"field": "options"})

    max_depth_raw = options.get("maxDepth", 8)
    try:
        max_depth = int(max_depth_raw)
    except (TypeError, ValueError):
        return _api_v1_error("VALIDATION_ERROR", "Поле options.maxDepth должно быть числом.", details={"field": "options.maxDepth"})
    if max_depth > 8:
        return _api_v1_error(
            "VALIDATION_ERROR",
            "Максимальная глубина prefetch_graph ограничена 8.",
            details={"field": "options.maxDepth", "max_allowed": 8},
        )
    options["maxDepth"] = max_depth

    include_trace = bool(options.get("includeTrace", False))
    resolver = TokenResolverService(logger=logger)
    try:
        result = resolver.prefetch_graph(
            document_id=document_id,
            context=context,
            tokens=tokens_list,
            options=options,
            include_trace=include_trace,
        )
    except TokenResolveError as exc:
        return _api_v1_error(exc.code, exc.message, details=exc.details)
    except Exception as exc:
        logger.exception("prefetch_graph API failed for document %s", document_id)
        return _api_v1_error(
            "SERVER_ERROR",
            "Не удалось выполнить prefetch графа токенов.",
            status=500,
            details={"exc": str(exc)},
        )
    return JsonResponse(result)



@csrf_exempt
def AddDocumentToUser(request, id=None):
    """Копирование документа для пользователя (API/внутреннее использование)."""
    if id is not None:
        document_id = id
    else:
        try:
            if request.body:
                try:
                    payload = json.loads(request.body.decode('utf-8'))
                except json.JSONDecodeError:
                    payload = request.POST
            else:
                payload = request.POST
            document_id = payload.get('id')
        except Exception as exc:
            message = f"Некорректные данные запроса: {exc}"
            logger.error(message)
            return JsonResponse({'success': False, 'error': message}, status=400)

    if document_id is None:
        message = "Не указан идентификатор документа."
        logger.error(message)
        return JsonResponse({'success': False, 'error': message}, status=400)

    success, document, error = _copy_document_for_user(request, document_id)
    if success:
        return JsonResponse({'success': True, 'id': document.id})

    message = f"Не удалось добавить документ пользователю: {error}"
    logger.error(message)
    return JsonResponse({'success': False, 'error': message}, status=400)

@csrf_exempt
def SaveSavedElement(request):
    """Обработчик запроса для сохранения сохраненного элемента"""
    response = HttpResponse()
    request_data = request.body
    stroke = json.loads(request_data)
    if stroke['id'] == '-1':
        saved_element = SavedElements()
    else:
        saved_element = SavedElements.objects.filter(id=stroke['id'])[0]
    saved_element.name = stroke['name']
    saved_element.json = stroke['json']
    saved_element.save()
    response['id'] = saved_element.id
    return response

@csrf_exempt
def DeleteDocument(request):
    """Обработчик запроса для удаления документа из базы данных"""
    response = redirect('/')
    request_data = request.body
    stroke = json.loads(request_data) 
    print(stroke)
    if stroke['id'] == '-1':
        return response
    document = DocumentsPattern.objects.get(id=stroke['id'])
    print(document)
    childs = Document_ParentDocument.objects.filter(parent=document) | Document_ParentDocument.objects.filter(document=document)
    print(childs)
    for child in childs:
        child.delete()
    try:
        file_path = Path(settings.MEDIA_ROOT) / str(document.file)
        if file_path.exists():
            file_path.unlink()
        picture_path = Path(settings.MEDIA_ROOT) / str(document.picture)
        if picture_path.exists():
            picture_path.unlink()
        document.delete()
        return response
    except Exception as exc:
        print(document)
        try:
            doc_deleted = document.delete()
        except Exception as exc:
            print(exc)
        print("#######", doc_deleted)
        print(exc)
        return response
    
def replace_last(source_string, replace_what, replace_with):
    head, _sep, tail = source_string.rpartition(replace_what)
    return head + replace_with + tail

@login_required
def ViewDocument(request):
    """Обработчик запроса для просмотра документа"""
    fileid = request.GET.get('id')
    Doc = get_object_or_404(DocumentsPattern, pk=fileid)
    file_path = Doc.file
    file_path = os.path.join(settings.MEDIA_ROOT, str(file_path))
    if request.user != Doc.owner.user:
        success, cloned_doc, error = _copy_document_for_user(request, fileid)
        if not success:
            message = f"Не удалось добавить документ пользователю: {error}"
            logger.warning(message)
            return HttpResponseBadRequest(message)
        return redirect(f'/document/view?id={cloned_doc.id}')
    parent_document = Document_ParentDocument.objects.filter(document=Doc.id)
    hide_legacy_linked_params = bool(getattr(settings, "DBM_DISABLE_LEGACY_LINKED_PARAMS", False))
    objects = []
    document_object_rows = list(
        DocumentPattern_Objects.objects.filter(document=Doc.id).select_related("object")
    )
    for doc_object in document_object_rows:
        base_object = doc_object.object
        base_params = list(sorted(Parameter.objects.filter(object=base_object), key=lambda x: x.id))
        if hide_legacy_linked_params:
            base_params = [
                parameter
                for parameter in base_params
                if not (
                    getattr(parameter, "linked_object_id", None)
                    and not getattr(parameter, "link_meta_id", None)
                    and not bool(getattr(parameter, "is_managed_link_param", False))
                )
            ]
        params = list(base_params)
        for param in base_params:
            if not hasattr(param, "linked_object") or not param.linked_object:
                continue
            relation = Object_ParentObject.objects.filter(
                parent_object=base_object,
                object=param.linked_object,
            ).first()
            if relation is None:
                continue
            link_metas = list(
                ObjectLinkMeta.objects.filter(
                    parent_object=base_object,
                    child_object=param.linked_object,
                ).order_by("order", "id")
            )
            if not link_metas:
                link_metas = [None]
            linked_params = list(Parameter.objects.filter(object=param.linked_object).order_by("id"))

            class PseudoParam:
                def __init__(
                    self,
                    original,
                    prefix,
                    linked_obj,
                    *,
                    link_type_value,
                    parent_id=None,
                    link_meta_id=None,
                ):
                    self.id = original.id
                    self.name = f"{prefix}.{original.name}"
                    self.identificator = False
                    if link_type_value == "single":
                        self.data_type = original.data_type
                    elif link_type_value == "multiple":
                        self.data_type = "ARRAY"
                    else:
                        self.data_type = original.data_type
                    self.linked_object = linked_obj
                    self.isChild = True
                    self.parent_id = parent_id
                    self.link_meta_id = link_meta_id

            for link_meta in link_metas:
                prefix = (link_meta.display_name if link_meta is not None else param.name) or param.name
                link_type_value = (link_meta.link_type if link_meta is not None else relation.link_type) or relation.link_type
                for linked_param in linked_params:
                    pseudo = PseudoParam(
                        linked_param,
                        prefix,
                        param.linked_object,
                        link_type_value=link_type_value,
                        parent_id=param.id,
                        link_meta_id=(link_meta.id if link_meta is not None else None),
                    )
                    params.append(pseudo)
        objects.append({"object": base_object, "params": params})

    doc_token_index = build_doc_token_index(Doc.id)
    context = {
        'title': 'Просмотр документа',
        'Doc': Doc,
        'fonts': Fonts.objects.order_by('id'),
        'variable': json.dumps({var.variable.id: f"{var.variable.name}:{var.variable.meaning}" for var in Document_VariableBlock.objects.filter(document=Doc.id)}),
        'document_json': json.dumps(Doc.json),
        'objects': objects,
        'doc_token_index_json': json.dumps(doc_token_index, ensure_ascii=False),
    }
    if request.user.is_authenticated:
        context['profile'] = Profile.for_user(request.user)

    # добавляем идентификатор шаблона и (опционально) выбранные записи для панели
    context['pattern'] = Doc
    context['selected_ids_json'] = json.dumps({})
    return render(request, 'document/document_view_v3.html', context)

def CreateDocumentMultiple(request):
    if request.method == 'GET':
        documents = request.POST.getlist('object_documents[]')
        idents = request.POST.getlist('object_idents[]')
        for doc_id in documents:
            return redirect(f'/document/view?id={res["id"]}')

def ViewDocumentAndCreateDocument(request):
    """Обработчик запроса для просмотра документа"""
    fileid = request.GET.get('id')
    
    Doc = get_object_or_404(DocumentsPattern, pk=fileid)
    file_path = Doc.file
    file_path = os.path.join(settings.MEDIA_ROOT, str(file_path))
    if request.user != Doc.owner.user:
        success, cloned_doc, error = _copy_document_for_user(request, fileid)
        if not success:
            message = f"Не удалось добавить документ пользователю: {error}"
            logger.warning(message)
            return HttpResponseBadRequest(message)
        return redirect(f'/document/view?id={cloned_doc.id}')
    parent_document = Document_ParentDocument.objects.filter(document=Doc.id)
    hide_legacy_linked_params = bool(getattr(settings, "DBM_DISABLE_LEGACY_LINKED_PARAMS", False))
    objects = []
    for doc_object in DocumentPattern_Objects.objects.filter(document=Doc.id):
        params = list(Parameter.objects.filter(object=doc_object.object))
        if hide_legacy_linked_params:
            params = [
                parameter
                for parameter in params
                if not (
                    getattr(parameter, "linked_object_id", None)
                    and not getattr(parameter, "link_meta_id", None)
                    and not bool(getattr(parameter, "is_managed_link_param", False))
                )
            ]
        objects.append({'object': doc_object.object, 'params': params})
    print(objects)
    doc_token_index = build_doc_token_index(Doc.id)
    context = {
        'title': 'Просмотр документа',
        'Doc': Doc,
        'fonts': Fonts.objects.order_by('id'),
        'variable': json.dumps({var.variable.id: f"{var.variable.name}:{var.variable.meaning}" for var in Document_VariableBlock.objects.filter(document=Doc.id)}),
        'document_json': json.dumps(Doc.json),
        'objects': objects,
        'create_document': True,
        'doc_token_index_json': json.dumps(doc_token_index, ensure_ascii=False),
    }
    if request.user.is_authenticated:
        context['profile'] = Profile.for_user(request.user)
    context['pattern'] = Doc
    context['selected_ids_json'] = json.dumps({})
    return render(request, 'document/document_view_v3.html', context)

def SavedElementFromJSON(json_stroke):
    """Функция создаёт элементы класса SavedElements из JSON-строки"""
    for list_elem in json_stroke:
        name, id_elem, element = list_elem['name'], list_elem['id'], list_elem['json']
        if id_elem == '-1':
            saved = SavedElements()
            saved.name = name
            saved.json = json.loads(element)
            saved.save()
        else:
            saved = SavedElements.objects.get(id=id_elem)
            saved.name = name
            saved.json = json.loads(element)
            saved.save()

@csrf_exempt
def UpdateDocument(request):
    """Функция обрабатывает запрос и обновляет документ в базе данных"""
    fileid = request.GET.get('id')
    doc = DocumentsPattern.objects.get(id=fileid)
    childs = Document_ParentDocument.objects.filter(parent=doc)
    print(childs)
    for child in childs:
        print(child.document.id)
    file_path = os.path.join(settings.MEDIA_ROOT, str(doc.file))
    document = Document()
    request_data = request.body
    stroke = json.loads(request_data)
    doc.json = stroke
    doc.name = stroke['doc_name']
    if not doc.save():
        response = HttpResponseNotModified()
        response['MessageOfError'] = "Документ не сохранён. Попробуйте позже.".encode('utf-8')
        response['TypeOfError'] = "Документ не сохранён".encode('utf-8')
        return response
    document.from_json(stroke)
    document.save(file_path)
    response = HttpResponse()
    return response

@csrf_exempt
def AcceptFilters(request):
    request_data = request.body
    try:
        stroke = json.loads(request_data)
    except Exception:
        stroke = {}
    if not isinstance(stroke, dict):
        stroke = {}
    result = {}
    for key, value in stroke.items():
        if not isinstance(value, dict):
            result[key] = ""
            continue

        phrase = value.get('phrase')
        if phrase is None:
            phrase = value.get('data-invis')
        if phrase is None:
            phrase = value.get('textContent')
        payload = dict(value)
        payload['phrase'] = '' if phrase is None else str(phrase)

        filters_raw = value.get('filters')
        if isinstance(filters_raw, dict):
            filters_iter = list(filters_raw.keys())
        elif isinstance(filters_raw, (list, tuple, set)):
            filters_iter = list(filters_raw)
        else:
            filters_iter = []

        for filter in filters_iter:
            if filter == 'Raskrit':
                payload['phrase'] = Raskritie(payload)
            if filter == 'ChangeCattle':
                payload['phrase'] = ChangeCattle(payload)
            if filter == 'ChangeCase':
                payload['phrase'] = UpperCase(payload)
        result[key] = str(payload.get('phrase', '')).replace('ё', 'е').replace('Ё', 'Е')
    response = HttpResponse()
    response.content = json.dumps(result)
    response.charset = 'utf-8'
    return response

def Raskritie(stroke):
    phrase = str(stroke['phrase'])
    newphrase = phrase
    for key, value in raskrit.items():
        newphrase = newphrase.replace(key, value)  
    return newphrase

def ItFIO(phrase):
    splitted = phrase.split()
    if len(splitted) != 3:
        return False
    res = 0
    for word in splitted:
        parsed_word = morph.parse(word)[0]
        print(morph.parse(word))
        if 'Name' in parsed_word.tag or 'Surn' in parsed_word.tag or 'Patr' in parsed_word.tag:
            res += 1
    return res >= 2

def GetGenderFIO(phrase):
    res = {'femn': 0, 'masc': 0, 'neut': 0}
    for word in phrase.split():
        parsed_word = morph.parse(word)[0]
        if parsed_word.tag.gender:
            res[parsed_word.tag.gender] += 1
    res_sorted = sorted(res, key=lambda x: res[x], reverse=True)
    return res_sorted[0]

def remove_quoted_text(phrase):
    pattern = r'(["«].*?["»])'
    match = re.search(pattern, phrase, re.DOTALL)
    if match:
        return True, match.group(0)
    return False, ''

def ChangeCattle(stroke):
    phrase = str(stroke['phrase'])
    print(phrase)
    ret_quotes = {}
    have, found = remove_quoted_text(phrase)
    while have:
        phrase = phrase.replace(found, f'(*{len(ret_quotes)}*)')
        ret_quotes[len(ret_quotes)] = found
        have, found = remove_quoted_text(phrase)
    if ItFIO(phrase):
        gend = GetGenderFIO(phrase)
        print(gend)
        params = {gend, stroke['filters']['ChangeCattle']}
    else:
        params = {stroke['filters']['ChangeCattle']}
    
    newphrase = ''
    for key, value in matchers.items():
        result = re.search(key, phrase, re.DOTALL)
        if result:
            phrase1 = result.group(0)
            try:
                splitted_phrase = re.split(value, phrase1)[1]
                phrase = phrase.replace(splitted_phrase, f'(*{len(ret_quotes)}*)')
                ret_quotes[len(ret_quotes)] = splitted_phrase
            except IndexError:
                splitted_phrase = ''
        else:
            continue
    for word in phrase.split():
        try:
            parsed_word = morph.parse(word)[0]
            newphrase += restore_capitalization(parsed_word.inflect(params).word, word)+ ' '
        except:
            newphrase += word+ ' '
            continue
    for key, value in ret_quotes.items():
        newphrase = newphrase.replace(f'(*{key}*)', value)
    return newphrase.strip()

def UpperCase(stroke):
    response = HttpResponse()
    phrase = str(stroke['phrase'])
    newphrase = phrase
    if stroke['filters']['ChangeCase'] == 'Upper':
        newphrase = phrase.upper()
    elif stroke['filters']['ChangeCase'] == 'Lower':
        newphrase = phrase.lower()
    return newphrase
    
def connect_objects_to_document(request, pk):
    if request.method == 'POST':
        print(request.POST)
        # Получить список выбранных объектов из запроса
        selected_objects = request.POST.getlist('selectedObjects[]')
        
        print(selected_objects)

        # Получить документ по идентификатору
        document = DocumentsPattern.objects.get(pk=pk)

        # Подключить выбранные объекты к документу
        for obj_id in selected_objects:
            obj = Object.objects.get(id=obj_id)
            doc_obj = DocumentPattern_Objects.objects.filter(document=document, object=obj).first()
            if doc_obj is None:
                doc_obj = DocumentPattern_Objects()
            doc_obj.document = document
            doc_obj.object = obj
            doc_obj.save()

        # Вернуть успешный ответ
        return HttpResponse()
    
def delete_object_from_document(request, pk):
    if request.method == 'POST':
        try:
            print(request.POST)
            # Получить список выбранных объектов из запроса
            object_id = request.POST['object_id']
            
            print(object_id)

            # Получить документ по идентификатору
            document = DocumentsPattern.objects.get(pk=pk)

            # Подключить выбранные объекты к документу
            doc_obj = DocumentPattern_Objects.objects.filter(document=document, object=object_id)
            doc_obj.delete()

            # Вернуть успешный ответ
            return HttpResponse()
        except Exception as e:
            print(e)
            return HttpResponseNotModified('Не удалено. Ошибка удаления. ' + e)
        
     
