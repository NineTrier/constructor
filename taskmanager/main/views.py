import json

from django.shortcuts import render, redirect
from django.views.generic import CreateView
from .models import Documents, Fonts, SavedElements, VariableBlock
from .forms import DocumentForm
import os
from django.conf import settings
from django.http import HttpResponse, Http404
from docx import Document
from .Document import Document

from .Formatter.Formatter import Formatter
from django.views.decorators.csrf import csrf_exempt
from transliterate import translit
from transliterate.decorators import transliterate_function


# Класс, который помогает создавать новый записи в базу данных Documents
# и открывает страницу с добавлением новых документов
class DocumentCreate(CreateView):
    model = Documents
    form_class = DocumentForm

    extra_context = {'documents': Documents.objects.all()}

    template_name = 'main/document_create.html'

    def get_success_url(self):
        return f"/document/view?id={self.object.id}&type=1"

def translit_russian(text):
    try:
        translited = translit(text, reversed=True)
        return translited
    except:
        return text
    

def create_New_Document(request):
    if request.method == 'POST':
        document = Documents()
        document.name = request.POST['name']
        document.description = request.POST['description']
        document.owner = request.POST['owner']
        document.file = f"documents\\{translit_russian(document.name)}{translit_russian(document.owner)}.docx"
        file_path = os.path.join(settings.MEDIA_ROOT, str(document.file))
        file_path_blank = os.path.join(settings.MEDIA_ROOT, f"documents\\blank.docx")
        doc_file = Document(file_path_blank)
        doc_file.save(file_path)
        document.save()
        print(document.id)
        return redirect(f"/document/view?id={document.id}&type=1")
    form = DocumentForm()
    context = {
        'form': form,
    }
    return render(request, 'main/new_document.html', context)


def uploadRed(request):
    filepath = request.GET.get('filename')
    filepath = os.path.normpath(filepath)
    print(filepath)
    name = "".join(filepath.split('\\')[-1].split('.')[:-1])
    print(name)
    file_path = os.path.join(settings.MEDIA_ROOT, str(filepath.split('\\')[-1]))
    doc = Documents(name=name, owner="Александр", description="Описание", file=file_path)
    doc.save()
    docToSave = Document(filepath)
    docToSave.save(file_path)
    file_path = os.path.join(settings.MEDIA_ROOT, str(filepath))
    if os.path.exists(file_path):
        frm = Formatter(file_path, path_to_save=str(settings.MEDIA_ROOT) + '/documents/Редактированные')
        frm.Redact()
        if os.path.exists(frm.path):
            with open(frm.path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
                response['Content-Disposition'] = 'inline; filename=' + os.path.basename(frm.path)
                return response
    raise Http404


# Функция, которая позволяет скачать файл, загруженный на сервер
def download(request):
    filepath = Documents.objects.filter(id=request.GET.get('id'))[0].file
    file_path = os.path.join(settings.MEDIA_ROOT, str(filepath))
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
            return response
    raise Http404


# Функция позволяет скачать изменённый при помощи летней программы файл
def change_doc(request):
    filepath = Documents.objects.filter(name=request.GET.get('filename'))[0].file
    file_path = os.path.join(settings.MEDIA_ROOT, str(filepath))
    if os.path.exists(file_path):
        frm = Formatter(file_path, path_to_save=str(settings.MEDIA_ROOT)+'/documents/Редактированные')
        frm.Redact()
        if os.path.exists(frm.path):
            with open(frm.path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
                response['Content-Disposition'] = 'inline; filename=' + os.path.basename(frm.path)
                return response
    raise Http404

@csrf_exempt
def DeleteSavedElement(request):
    response = redirect('/')
    request_data = request.body
    stroke = json.loads(request_data)
    if stroke['id'] == '-1':
        return response
    SavedElements.objects.filter(id=stroke['id']).delete()
    return response

@csrf_exempt
def DeleteVariable(request):
    response = redirect('/')
    request_data = request.body
    stroke = json.loads(request_data)
    if stroke['id'] == '-1':
        return response
    VariableBlock.objects.filter(id=stroke['id']).delete()
    return response

@csrf_exempt
def DeleteDocument(request):
    response = redirect('/')
    request_data = request.body
    stroke = json.loads(request_data)
    if stroke['id'] == '-1':
        return response
    Documents.objects.filter(id=stroke['id']).delete()
    return response

def ViewDocument(request):
    fileid = request.GET.get('id')
    print(request.GET.get('type'))
    Doc = Documents.objects.filter(id=fileid)[0]
    file_path = Doc.file
    print(file_path)
    file_path = os.path.join(settings.MEDIA_ROOT, str(file_path))
    document = Document(file_path)
    context = {
        'title': 'Просмотр документа',
        'id_doc': fileid,
        'document': [(child, child.id) for child in document.childs],
        'name_of_document': Doc.name,
        'sectPr': document.childs[-1],
        'fonts': Fonts.objects.order_by('id'),
        'saved_elements': json.dumps({f"{el.name}:{el.id}": json.dumps(el.json) for el in SavedElements.objects.filter(type=Doc.type)}),
        'variable': json.dumps({var.id: f"{var.name}:{var.meaning}" for var in VariableBlock.objects.filter(doc=Doc.id)}),
        'document_json': json.dumps(Doc.json),
        'type': request.GET.get('type')
    }
    return render(request, 'main/document_view.html', context)


def SavedElementFromJSON(json_stroke):
    for list_elem in json_stroke:
        print(list_elem)
        name, id_elem, element = list_elem['name'], list_elem['id'], list_elem['json']
        if id_elem == '-1':
            saved = SavedElements()
            saved.name = name
            saved.json = json.loads(element)
            saved.save()
        else:
            saved = SavedElements.objects.get(id=id_elem)
            saved.name = name
            saved.json = json.loads(element)
            saved.save()

def VariablesFromJSON(json_stroke, id_doc):
    for variable in json_stroke:
        print(variable)
        name, id_elem, value = variable['name'], variable['id'], variable['value']
        if id_elem == '-1':
            var = VariableBlock()
            var.name = name
            var.meaning = value
            var.doc_id = id_doc
            var.save()
        else:
            var = VariableBlock.objects.get(id=id_elem)
            var.name = name
            var.meaning = value
            var.save()


@csrf_exempt
def UpdateDocument(request):
    fileid = request.GET.get('id')
    doc = Documents.objects.filter(id=fileid)[0]
    file_path = os.path.join(settings.MEDIA_ROOT, str(doc.file))
    document = Document(file_path)
    request_data = request.body
    stroke = json.loads(request_data)
    doc.json = stroke
    doc.name = stroke['doc_name']
    doc.save()
    print(stroke)
    SavedElementFromJSON(stroke["list_saved"])
    VariablesFromJSON(stroke["variables"], fileid)
    document.from_json(stroke)
    document.save(file_path)
    response = redirect('/')
    return response


# Главная страница веб-сервиса
def index(request):
    documents = Documents.objects.order_by('-id')
    context = {
        'title': 'Главная страница сайта',
        'documents': documents,
    }
    return render(request, 'main/index.html', context)


# Страница about
def about(request):
    return render(request, 'main/about.html')
