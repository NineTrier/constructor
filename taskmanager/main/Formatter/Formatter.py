import os

from docx import Document
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

allowed = ['.', ',', '-', ':', ';', '"', "'"]

default = f"{os.path.join(os.path.join(os.environ['USERPROFILE']), 'Documents')}\\Ассистент"
flag = False


def get_all_Date_Time(text):
    dates = []
    times = []
    textSplitted = str(text).split(' ')
    for i in textSplitted:
        try:
            boolDat, Dat = is_Date(i if 8 <= len(i) <= 15 else "nothing")
            boolTime, Time = is_Time(i if 5 <= len(i) <= 8 else "nothing")
        except:
            continue
        if boolDat:
            dates.append(Dat)
        if boolTime:
            times.append(Time)
    return dates, times


def get_textInput(self, paragraph):
    run = paragraph.add_run()
    self.create_attribute(run._r, 'w:rsidRPr', '00921D4A')
    rPr = self.create_element('w:rPr')
    rPr1 = self.create_element('w:szCs')
    self.create_attribute(rPr1, 'w:val', '26')
    rPr2 = self.create_element('w:highlight')
    self.create_attribute(rPr2, 'w:val', 'default')
    rPr.append(rPr1)
    rPr.append(rPr2)
    run._r.append(rPr)

    fldStart = self.create_element('w:fldChar')
    self.create_attribute(fldStart, 'w:fldCharType', 'begin')
    ffdata = self.create_element('w:ffData')
    name = self.create_element('w:name')
    self.create_attribute(name, 'w:val', 'Тратата')
    ffdata.append(name)
    enabled = self.create_element('w:enabled')
    ffdata.append(enabled)
    calc = self.create_element('w:calcOnExit')
    self.create_attribute(calc, 'w:val', '0')
    ffdata.append(calc)
    textInput = self.create_element('w:textInput')
    default = self.create_element('w:default')
    self.create_attribute(default, 'w:val', "Текст")
    textInput.append(default)
    ffdata.append(textInput)
    fldStart.append(ffdata)
    run._r.append(fldStart)

    run2 = paragraph.add_run()
    self.create_attribute(run2._r, 'w:rsidRPr', '00921D4A')
    rPrN = self.create_element('w:rPr')
    rPrN1 = self.create_element('w:szCs')
    self.create_attribute(rPr1, 'w:val', '26')
    rPrN.append(rPrN1)
    run._r.append(rPrN)
    instrText = self.create_element('w:instrText')
    self.create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = " FORMTEXT "
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    self.create_attribute(run3._r, 'w:rsidRPr', '00921D4A')
    rPrNN = self.create_element('w:rPr')
    fldChar1 = self.create_element('w:szCs')
    self.create_attribute(fldChar1, 'w:val', '26')
    rPrNN.append(fldChar1)
    run3._r.append(rPrNN)

    run4 = paragraph.add_run()
    self.create_attribute(run4._r, 'w:rsidRPr', '00921D4A')
    rPrNNN = self.create_element('w:rPr')
    fldChar2 = self.create_element('w:szCs')
    self.create_attribute(fldChar2, 'w:val', '26')
    rPrNNN.append(fldChar2)
    run4._r.append(rPrNNN)
    fldCharSep = self.create_element('w:fldChar')
    self.create_attribute(fldCharSep, 'w:fldCharType', 'separate')
    run4._r.append(fldCharSep)

    run5 = paragraph.add_run()
    self.create_attribute(run5._r, 'w:rsidRPr', '00921D4A')
    rPrNNNN = self.create_element('w:rPr')
    fldChar22 = self.create_element('w:szCs')
    self.create_attribute(fldChar22, 'w:val', '26')
    rPrNNNN.append(fldChar22)
    run5._r.append(rPrNNNN)
    fldCharText = self.create_element('w:t')
    fldCharText.text = "Текст"
    run5._r.append(fldCharText)

    run6 = paragraph.add_run()
    self.create_attribute(run6._r, 'w:rsidRPr', '00921D4A')
    fldEnd = self.create_element('w:fldChar')
    self.create_attribute(fldEnd, 'w:fldCharType', 'end')
    run6._r.append(fldEnd)


def is_Date(text: str):
    if text[2] == '.' or text[2] == ',' and text[5] == '.' or text[5] == ',':
        for i in range(10):
            if i == 2 or i == 5:
                continue
            if not ('0' <= text[i] <= '9'):
                if i == 8 and (text[9] == " " or text[9:11] == 'г.' or text[9] in allowed):
                    return True, f"{text[:8]}"
                return False, ""
        return True, text[:10]
    else:
        return False, ""


def is_Time(text: str):
    if text[2] == ':':
        for i in range(len(text)):
            if i == 2:
                continue
            if not ('0' <= text[i] <= '9'):
                return False, ""
        return True, text[:5]
    else:
        return False, ""


class Formatter:
    settings: dict

    month = {"01": "января", "02": "февраля", "03": "марта", "04": "апреля", "05": "мая",
             "06": "июня", "07": "июля", "08": "августа", "09": "сентября", "10": "октября",
             "11": "ноября", "12": "декабря"}

    abr = [(' РС (Я) ', ' Республики Саха (Якутия) '), (' РБ ', ' Республики Бурятия '),
           (' ИО ', ' Иркутской области '), (' ЗК', ' Забайкальского края ')]

    pad = {'Арбитражный суд/Арбитражного суда': ['ение', 'содействии', 'Содействии'],
           'Арбитражный суд/Арбитражном суде': ['наличие в', 'Наличие в'],
           'Арбитражный суд/Арбитражному суду': ['поручил', 'Поручил']}

    def __init__(self, path, path_to_save, settings=""):
        self.doc = Document(path)
        self.path = path
        self.name = ""
        self.number = ""
        self.settings = {'ChangeNumber': True, 'ChangeDate': True, 'ChangeKavich': True, 'ChangeTN': True,
                         'ChangeTire': True, 'ChangePadeg': True, 'ChangeRF': False, 'ChangeGod': True,
                         'ChangeTime': True, 'ChangeHighlight': False, 'ChangeColor': True, 'DeleteUnder': True,
                         'PathToSave': default, 'FontFamily': "Times New Roman", 'FontSize': 12,
                         'Interval': 1.5}
        self.path_to_save = path_to_save
        self.path_to_save_dif = f"{self.path_to_save}\\Разное"
        self.Dates = []
        self.Times = []

    def create_element(self, name):
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(self, run):
        fldStart = self.create_element('w:fldChar')
        self.create_attribute(fldStart, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'separate')

        fldChar2 = self.create_element('w:t')
        fldChar2.text = "2"

        fldEnd = self.create_element('w:fldChar')
        self.create_attribute(fldEnd, 'w:fldCharType', 'end')

        run._r.append(fldStart)
        run._r.append(instrText)
        run._r.append(fldChar1)
        run._r.append(fldChar2)
        run._r.append(fldEnd)

    def u_null(self, doc):
        try:
            run = doc.paragraphs[0].add_run()
            flag = False
            for bad in run._r.xpath("//w:rPr"):
                for i in bad:
                    if i.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}u':
                        i.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = 'none'
                        flag = True
                if not flag:
                    u = self.create_element('w:u')
                    self.create_attribute(u, 'w:val', 'none')
                    bad.append(u)
                flag = False
        except Exception as exc:
            print(exc, "u_null")

    def i_null(self, doc):
        try:
            run = doc.paragraphs[0].add_run()
            for bad in run._r.xpath("//w:rPr"):
                for i in bad:
                    if i.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i':
                        bad.remove(i)
        except Exception as exc:
            print(exc, "i_null")

    def b_null(self, doc):
        try:
            run = doc.paragraphs[0].add_run()
            for bad in run._r.xpath("//w:rPr"):
                for i in bad:
                    if i.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b':
                        bad.remove(i)
        except Exception as exc:
            print(exc, "b_null")

    def reformat_number(self, doc):
        run = doc.paragraphs[0].add_run()
        num_list = ['А19', 'А58', 'А78', 'А10', 'A19', 'A58', 'A78', 'A10']
        for j in num_list:
            try:
                for bad in run._r.xpath(f"//w:t[contains(text(),'{j}')]"):
                    for i in range(10, 99):
                        badSplitted = bad.text.split(f'/{i}')
                        if len(badSplitted) > 1:
                            if badSplitted[1] == '':
                                bad.text = str(bad.text).replace(f'/{i}', f'/20{i}')
                            elif '0' <= badSplitted[1][0] <= '9':
                                continue
                            else:
                                bad.text = str(bad.text).replace(f'/{i}', f'/20{i}')
                        else:
                            continue
            except Exception as exc:
                print(exc)
                continue

    def delete_from_lxml(self, doc, elem):
        try:
            run = doc.paragraphs[0].add_run()
            for bad in run._r.xpath(elem):
                bad.getparent().remove(bad)
        except Exception as exc:
            print(exc, "delete_from_lxml")

    def change_color_font(self, doc):
        try:
            run = doc.paragraphs[0].add_run()
            flag = False
            for bad in run._r.xpath("//w:rPr"):
                for i in bad:
                    if i.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color':
                        i.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = '000000'
                        i.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor'] = 'text1'
                        flag = True
                if not flag:
                    col = self.create_element('w:color')
                    self.create_attribute(col, 'w:val', '000000')
                    self.create_attribute(col, 'w:themeColor', 'text1')
                    bad.append(col)
                flag = False
        except Exception as exc:
            print(exc, "u_null")

    def find_hyperlinks(self):
        print('Ищу ссылки...')
        self.hyperlinks = []
        try:
            run = self.doc.paragraphs[0].add_run()
            for bad in run._r.xpath('//w:hyperlink'):
                print(bad.attrib, 'hyperlink')
                r = self.create_element('w:r')
                t = self.create_element('w:t')
                if '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id' in dict(bad.attrib).keys():
                    self.hyperlinks.append(
                        (bad,
                         f";{bad.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id']}"))
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = f";{bad.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id']}"
                else:
                    idh = len(self.hyperlinks)
                    self.hyperlinks.append(
                        (bad, f";hyper{idh}"))
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = f";hyper{idh}"
                r.append(t)
                parent = bad.getparent()
                bad.addnext(r)
                parent.remove(bad)
            print('Нашёл ссылки!')
            print(self.hyperlinks)
        except Exception as exc:
            print(exc, "find_hyperlinks")

    def revive_hyperlinks(self, doc):
        print('Восстанавливаю ссылки...')
        for h in self.hyperlinks:
            try:
                run = doc.paragraphs[0].add_run()
                bad = run._r.xpath(f"//w:t[contains(text(),'{h[1]}')]")[0]
                text = str(bad.text)
                textSplitted = text.split(h[1])
                parent = bad.getparent()
                if textSplitted[0] == '' and textSplitted.count('') == 1:
                    r = self.create_element('w:r')
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = textSplitted[1]
                    r.append(t)
                    bad.addnext(r)
                    bad.addnext(h[0])
                    parent.remove(bad)
                elif textSplitted[1] == '' and textSplitted.count('') == 1:
                    r = self.create_element('w:r')
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = textSplitted[0]
                    r.append(t)
                    bad.addnext(h[0])
                    bad.addnext(r)
                    parent.remove(bad)
                elif textSplitted.count('') > 1:
                    bad.addnext(h[0])
                    parent.remove(bad)
                elif textSplitted.count('') == 0:
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = textSplitted[0]
                    r = self.create_element('w:r')
                    r.append(t)
                    r1 = self.create_element('w:r')
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = textSplitted[1]
                    r1.append(t)
                    bad.addnext(r1)
                    bad.addnext(h[0])
                    bad.addnext(r)
                    parent.remove(bad)
                print('Восстановил ссылку!')
            except Exception as exc:
                print(exc, "revive_hyperlinks", h)
                continue

    def delete_textInput(self, paragraph):
        try:
            run = paragraph.add_run()
            for bad in run._r.xpath('//w:textInput'):
                bad.getparent().getparent().getparent().getparent().remove(bad.getparent().getparent().getparent())
            for bad in run._r.xpath("//w:instrText[text()=' FORMTEXT ']"):
                bad.getparent().getparent().remove(bad.getparent())
            for bad in run._r.xpath("//w:instrText[text()='FORMTEXT ']"):
                bad.getparent().getparent().remove(bad.getparent())
            for bad in run._r.xpath('//w:smartTag/w:r'):
                anc = bad.getparent()
                anc.addnext(bad)
                anc.getparent().remove(anc)
        except Exception as exc:
            print(exc, "delete_textInput")

    def find_highlight(self, paragraph, color='yellow'):
        massiv = []
        try:
            run = paragraph.add_run()
            for bad in run._r.xpath(f"//w:highlight[@w:val='{color}']"):
                for i in bad.getparent().getparent():
                    if i.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                        count = len(massiv)
                        i.text = self.zamena(i.text)
                        if f"{i.text};hg{count}" not in massiv:
                            massiv.append(
                                [f"{i.text};hg{count}", f";hg{count}"] if len(str(i.text).replace(' ', '')) > 1 else [
                                    f"_{i.text};hg{count}", f";hg{count}"])
                        i.text = f"{i.text};hg{count}" if len(
                            str(i.text).replace(' ', '')) > 1 else f"_{i.text};hg{count}"
            return massiv
        except Exception as exc:
            print(exc, "delete_textInput")

    def rewrite_highlights(self, hg, doc):
        for h in hg:
            run = doc.paragraphs[0].add_run()
            try:
                bad = run._r.xpath(f"//w:t[contains(text(),'{h[0]}')]")[0]
                text = str(bad.text)
                textSplitted = text.split(h[0])
                parent = bad.getparent()
                rPr = self.create_element('w:rPr')
                highlight = self.create_element('w:highlight')
                self.create_attribute(highlight, 'w:val', 'yellow')
                rPr.append(highlight)
                parent.remove(bad)
                if textSplitted[0] == '' and textSplitted.count('') == 1:
                    parent.append(rPr)
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = h[0].replace(h[1], '').replace('_', '')
                    parent.append(t)
                    r = self.create_element('w:r')
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = textSplitted[1]
                    r.append(t)
                    parent.addnext(r)
                elif textSplitted[1] == '' and textSplitted.count('') == 1:
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = textSplitted[0]
                    parent.append(t)
                    r = self.create_element('w:r')
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = h[0].replace(h[1], '').replace('_', '')
                    r.append(rPr)
                    r.append(t)
                    parent.addnext(r)
                elif textSplitted.count('') > 1:
                    parent.append(rPr)
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = h[0].replace(h[1], '').replace('_', '')
                    parent.append(t)
                elif textSplitted.count('') == 0:
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = textSplitted[0]
                    parent.append(t)
                    r = self.create_element('w:r')
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = h[0].replace(h[1], '').replace('_', '')
                    r.append(rPr)
                    r.append(t)
                    parent.addnext(r)
                    r1 = self.create_element('w:r')
                    t = self.create_element('w:t')
                    self.create_attribute(t, 'xml:space', 'preserve')
                    t.text = textSplitted[1]
                    r1.append(t)
                    r.addnext(r1)
            except Exception as exc:
                print(exc, "rewrite_highlights : ", h)
                continue

    def change_font(self, doc):
        run = doc.paragraphs[0].add_run()
        for bad in run._r.xpath(f"//w:rFonts"):
            try:
                parent = bad.getparent()
                rfont = self.create_element('w:rFonts')
                self.create_attribute(rfont, 'w:ascii', self.settings['FontFamily'])
                self.create_attribute(rfont, 'w:hAnsi', self.settings['FontFamily'])
                self.create_attribute(rfont, 'w:cs', self.settings['FontFamily'])
                self.create_attribute(rfont, 'w:eastAsia', self.settings['FontFamily'])
                parent.remove(bad)
                parent.append(rfont)
            except:
                continue
        for bad in run._r.xpath(f"//w:sz"):
            try:
                parent = bad.getparent()
                rfont = self.create_element('w:sz')
                self.create_attribute(rfont, 'w:val', str(self.settings['FontSize'] * 2))
                parent.remove(bad)
                parent.append(rfont)
            except:
                continue
        for bad in run._r.xpath(f"//w:szCs"):
            try:
                parent = bad.getparent()
                rfont = self.create_element('w:szCs')
                self.create_attribute(rfont, 'w:val', str(self.settings['FontSize'] * 2))
                parent.remove(bad)
                parent.append(rfont)
            except:
                continue

    def change_interval(self, doc):
        run = doc.paragraphs[0].add_run()
        flag = False
        for bad in run._r.xpath(f"//w:pPr"):
            try:
                for i in bad:
                    if i.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing':
                        i.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after'] = '0'
                        i.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line'] = str(
                            int(self.settings['Interval'] * 240))
                        i.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lineRule'] = 'auto'
                        flag = True
                if not flag:
                    rfont = self.create_element('w:spacing')
                    self.create_attribute(rfont, 'w:after', '0')
                    self.create_attribute(rfont, 'w:line', str(int(self.settings['Interval'] * 240)))
                    self.create_attribute(rfont, 'w:lineRule', 'auto')
                    bad.append(rfont)
                flag = False
            except Exception as exc:
                print(exc, 'change_interval')
                continue

    def zamena(self, text):
        # Убираем двойные пробелы
        if "  " in text:
            text = text.replace('  ', ' ')
        if ' т.е. ' in text:  # проверяем, если ли т.е., если есть заменяем на то есть
            text = text.replace(' т.е. ', ' то есть ')
        if list(self.settings.values())[0]:
            if " N " in text:  # проверяем, если ли N, если есть заменяем на №
                text = text.replace(' N ', ' № ')
        if list(self.settings.values())[2]:
            if '"' in text:  # проверяем, если ли "...", если есть заменяем на «...»
                text = text.replace(' "', ' «')
                text = text.replace('"', '»')
            if '“' in text and '”' in text:  # проверяем, если ли “...”, если есть заменяем на «...»
                text = text.replace('“', '«')
                text = text.replace('”', '»')
        if list(self.settings.values())[4]:
            if ' - ' in text:  # проверяем, если ли -, если есть заменяем на –
                text = text.replace('- ', '– ')
        # Раскрываем аббревиатуры
        if list(self.settings.values())[3]:
            for ab in self.abr:
                abFinded = text.find(ab[0])
                if abFinded != -1 and abFinded != text.find(f"{ab[0]} РФ"):
                    text = text[:abFinded] + ab[1] + text[abFinded + len(ab[0]):]
        # Работаем с падежами
        if list(self.settings.values())[5]:
            for pd in list(self.pad.keys()):
                pdSplit = pd.split('/')
                if text.find(pdSplit[0]) != -1:
                    poz = text.find(pdSplit[0])
                    for pds in self.pad[pd]:
                        poz1 = text.find(pds)
                        if poz1 != -1:
                            poz1 += len(pds)
                            if 0 <= poz - poz1 <= 8:
                                text = text.replace(pdSplit[0], pdSplit[1])
        # Отдельная опциональная аббревиатура
        if list(self.settings.values())[6]:
            if ' РФ ' in text:
                text = text.replace(' РФ ', ' Российской Федерации ')
        dates, times = get_all_Date_Time(text)
        self.Dates.append(dates)
        self.Times.append(times)
        if list(self.settings.values())[1]:
            if self.Dates[-1]:
                for date in self.Dates[-1]:
                    try:
                        text = text.replace(date, str(date).replace(',', '.'))
                        date = str(date).replace(',', '.')
                        splittedDate = date.split('.')
                        if len(splittedDate[2]) == 2:
                            text = text.replace(f"{date} года",
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                            text = text.replace(f"{date} г.",
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                            text = text.replace(f"{date}г.",
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                            text = text.replace(f"{date} г",
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                            text = text.replace(f"{date}г",
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                            text = text.replace(date,
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                        else:
                            text = text.replace(f"{date} года",
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                            text = text.replace(f"{date} г.",
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                            text = text.replace(f"{date}г.",
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                            text = text.replace(f"{date} г",
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                            text = text.replace(f"{date}г",
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                            text = text.replace(date,
                                                f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                    except Exception as exc:
                        print(f"Не смог поменять дату {date}", exc)
                        continue
        if self.Times[-1]:
            for timen in self.Times[-1]:
                splittedTime = timen.split(':')
                text = text.replace(timen, f"на {splittedTime[0]} часов {splittedTime[1]} минут")
        return text

    # нумерация
    def numbering(self):
        try:
            if self.doc.sections[0].header.paragraphs[0].text == "":
                self.add_page_number(self.doc.sections[0].header.paragraphs[0].add_run())
                self.doc.sections[0].header.paragraphs[
                    0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # выравниваем по центру
                self.doc.sections[
                    0].different_first_page_header_footer = True  # особый колонтитул для первой страницы - вкл
                sectPr = self.doc.sections[0]._sectPr
                pgNumType = OxmlElement('w:pgNumType')
                pgNumType.set(ns.qn('w:start'), "1")  # 1 это с какой страницы начинается отсчёт
                sectPr.append(pgNumType)
            return True
        except Exception as exc:
            print(exc, 'numbering')
            return False

    def check_folder_path(self):
        try:
            print(self.path, 1)
            file_name = os.path.basename(self.path)
            print(file_name, 2)
            number_doc = file_name[0:file_name.index("_", file_name.index("_") + 1)] if "_" in file_name else file_name
            print(file_name, 3)
            file_to_save = f"Отформатированный {file_name}"
            if not os.path.exists(self.path_to_save):
                print("Создаю ассистента")
                os.mkdir(self.path_to_save)
            user_file = str(self.path_to_save) + f"\\{number_doc}" if "_" in file_name else self.path_to_save_dif
            print(user_file, 4)
            if not os.path.exists(user_file):
                print("Создаю папку")
                os.mkdir(user_file)
            user_file += "\\" + file_to_save
            return user_file, number_doc, file_name
        except Exception as exc:
            print(exc, 'check_folder_path')

    # Функция для форматирования текст
    def Format(self):
        try:
            # Настройка отступов
            section = self.doc.sections[-1]
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)
            # section.left_margin = Mm(15)
            section.right_margin = Mm(15)
            section.header_distance = Mm(10)
            # отступ от нижнего края страницы до
            # нижнего края нижнего колонтитула
            section.footer_distance = Mm(10)

            print(self.doc.styles)
            for style in self.doc.styles:
                try:
                    style.font.name = self.settings['FontFamily']
                    style.font.size = Pt(self.settings['FontSize'])
                    # style.font.highlight_color
                except:
                    continue
            # Настройка междустрочного интервала и убираем выделение корректором

            for p in self.doc.paragraphs:
                for run in p.runs:
                    run.font.name = self.settings['FontFamily']
                    run.font.size = Pt(self.settings['FontSize'])
                # p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                # p.style.font.highlight_color = WD_COLOR_INDEX.AUTO
            return True
        except Exception as exc:
            print(exc, 'Format')
            return False

    # функция редактирование текста
    def Redact(self):
        print('Нумерую страницы...', end=" ")
        if self.numbering():
            print('Пронумеровал!')
        else:
            print('Не удалось пронумеровать :(')
        print('Форматирую текст...', end=" ")
        if self.Format():
            print('Отформатировал!')
        else:
            print('Не удалось отформатировать :(')

        self.find_hyperlinks()

        if not self.settings["ChangeHighlight"]:
            Highlights = self.find_highlight(self.doc.paragraphs[0])
        else:
            Highlights = []
        for hg in range(len(Highlights)):
            dat, tim = get_all_Date_Time(Highlights[hg])
            for d in dat:
                splittedDate = d.split('.')
                if len(splittedDate[2]) == 2:
                    Highlights[hg][0] = Highlights[hg][0].replace(f"{d} года",
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                    Highlights[hg][0] = Highlights[hg][0].replace(f"{d}г.",
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                    Highlights[hg][0] = Highlights[hg][0].replace(f"{d} г.",
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                    Highlights[hg][0] = Highlights[hg][0].replace(f"{d}г",
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                    Highlights[hg][0] = Highlights[hg][0].replace(f"{d} г",
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                    Highlights[hg][0] = Highlights[hg][0].replace(d,
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} 20{splittedDate[2]} года")
                else:
                    Highlights[hg][0] = Highlights[hg][0].replace(f"{d} года",
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                    Highlights[hg][0] = Highlights[hg][0].replace(f"{d}г.",
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                    Highlights[hg][0] = Highlights[hg][0].replace(f"{d} г.",
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                    Highlights[hg][0] = Highlights[hg][0].replace(f"{d}г",
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                    Highlights[hg][0] = Highlights[hg][0].replace(f"{d} г",
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
                    Highlights[hg][0] = Highlights[hg][0].replace(d,
                                                                  f"{splittedDate[0]} {self.month[splittedDate[1]]} {splittedDate[2]} года")
            for t in tim:
                splittedTime = t.split(':')
                Highlights[hg][0] = Highlights[hg][0].replace(t, f"{splittedTime[0]} часов {splittedTime[1]} минут")
            num_list = ['А19', 'А58', 'А78', 'А10', 'A19', 'A58', 'A78', 'A10']
            for j in num_list:
                if j in Highlights[hg]:
                    for i in range(10, 99):
                        if f"/{i}" in Highlights[hg]:
                            Highlights[hg][0] = Highlights[hg][0].replace(f"/{i}", f"/20{i}")
        print(Highlights)
        for p in self.doc.paragraphs:  # проходим все абзацы в документе на поиск ошибок, и заменяем их
            self.delete_textInput(p)
            text = self.zamena(str(p.text))
            if text == "":
                continue
            p.text = text
            # флаг, который отвечает, есть ли ошибка в абзаце
            # если есть, то правим и заменяем текс, если нет, то нет
        print(self.Dates)
        print(self.Times)
        path_file, number_file, name_file = self.check_folder_path()
        self.doc.save(path_file)

        doc = Document(path_file)
        if len(Highlights) > 0:
            if self.settings["ChangeHighlight"]:
                self.delete_from_lxml(doc, f"//w:highlight")
            else:
                self.rewrite_highlights(Highlights, doc)
        if list(self.settings.values())[7]:
            self.reformat_number(doc)
        self.revive_hyperlinks(doc)
        self.i_null(doc)
        self.b_null(doc)
        if self.settings['DeleteUnder']:
            self.u_null(doc)
        if self.settings['ChangeColor']:
            self.change_color_font(doc)
        self.change_font(doc)
        self.change_interval(doc)
        doc.save(path_file)

        # os.startfile(path_file)
        self.path = path_file
        self.number = number_file
        self.name = name_file
